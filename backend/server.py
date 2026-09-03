"""
Skrivestemme — Backend
Norsk skriveassistent som lærer brukerens stemme og hjelper med skrivesperre.
"""
import os
import io
import uuid
import logging
import re
import json
import asyncio
import difflib
from pathlib import Path
from datetime import datetime, timezone, timedelta
from typing import List, Optional, Literal
from urllib.parse import urlencode

import httpx
import bcrypt
import secrets
import requests
import stripe
from PIL import Image
from fastapi import FastAPI, APIRouter, HTTPException, Depends, Request, Response, UploadFile, File, Form, Query, Header
from fastapi.responses import StreamingResponse, JSONResponse, RedirectResponse
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
from pydantic import BaseModel, Field, ConfigDict

import litellm
from openai import AsyncOpenAI

# Extraction libs
from PyPDF2 import PdfReader
from docx import Document
from docx.shared import Pt, Cm
from docx import Document as DocxDocument


ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

# Mongo
mongo_url = os.environ['MONGO_URL']
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ['DB_NAME']]

EMERGENT_LLM_KEY = os.environ.get('EMERGENT_LLM_KEY')  # still used by storage_put/storage_get below
ANTHROPIC_API_KEY = os.environ.get('ANTHROPIC_API_KEY')
OPENAI_API_KEY = os.environ.get('OPENAI_API_KEY')
GEMINI_API_KEY = os.environ.get('GEMINI_API_KEY')
XAI_API_KEY = os.environ.get('XAI_API_KEY')
stripe.api_key = os.environ.get('STRIPE_SECRET_KEY')
STRIPE_WEBHOOK_SECRET = os.environ.get('STRIPE_WEBHOOK_SECRET', '')
BETA_FREE_SLOTS = 10
FOUNDER_SLOTS = 100  # First 100 (including the 50 beta) can access founder prices

# ---------- Object Storage ----------
STORAGE_URL = "https://integrations.emergentagent.com/objstore/api/v1/storage"
APP_NAME = "bragr"
_storage_key: Optional[str] = None


def init_storage() -> Optional[str]:
    global _storage_key
    if _storage_key:
        return _storage_key
    if not EMERGENT_LLM_KEY:
        return None
    try:
        resp = requests.post(f"{STORAGE_URL}/init", json={"emergent_key": EMERGENT_LLM_KEY}, timeout=30)
        resp.raise_for_status()
        _storage_key = resp.json()["storage_key"]
        return _storage_key
    except Exception as e:
        logging.getLogger(__name__).warning(f"Storage init failed: {e}")
        return None


def storage_put(path: str, data: bytes, content_type: str) -> dict:
    key = init_storage()
    if not key:
        raise HTTPException(status_code=500, detail="Fillagring ikke tilgjengelig")
    resp = requests.put(
        f"{STORAGE_URL}/objects/{path}",
        headers={"X-Storage-Key": key, "Content-Type": content_type},
        data=data, timeout=120,
    )
    if resp.status_code == 403:
        # key expired — reinit and retry once
        global _storage_key
        _storage_key = None
        key = init_storage()
        resp = requests.put(
            f"{STORAGE_URL}/objects/{path}",
            headers={"X-Storage-Key": key, "Content-Type": content_type},
            data=data, timeout=120,
        )
    resp.raise_for_status()
    return resp.json()


def storage_get(path: str) -> tuple[bytes, str]:
    key = init_storage()
    if not key:
        raise HTTPException(status_code=500, detail="Fillagring ikke tilgjengelig")
    resp = requests.get(
        f"{STORAGE_URL}/objects/{path}",
        headers={"X-Storage-Key": key}, timeout=60,
    )
    if resp.status_code == 403:
        global _storage_key
        _storage_key = None
        key = init_storage()
        resp = requests.get(
            f"{STORAGE_URL}/objects/{path}",
            headers={"X-Storage-Key": key}, timeout=60,
        )
    if resp.status_code == 404:
        raise HTTPException(status_code=404, detail="Fil ikke funnet i lagring")
    resp.raise_for_status()
    return resp.content, resp.headers.get("Content-Type", "application/octet-stream")




app = FastAPI(title="Bragarmål")
api_router = APIRouter(prefix="/api")


# ─── Health probes ──────────────────────────────────────────────
# Kubernetes / load balancer hits `GET /health` (no `/api` prefix).
# Also expose `/api/health` for consistency with the rest of the API surface.
@app.get("/health")
async def _health():
    return {"status": "ok"}


@api_router.get("/health")
async def _api_health():
    return {"status": "ok"}


@app.on_event("startup")
async def _startup():
    init_storage()

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)


# ---------- Models ----------
class User(BaseModel):
    model_config = ConfigDict(extra="ignore")
    user_id: str
    email: str
    name: str
    picture: Optional[str] = None
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))


SAMPLE_CATEGORIES = {
    "ren_menneske_gammel": "Ren menneske · gamle tekster",
    "ren_menneske_ny": "Ren menneske · nye tekster",
    "hybrid": "Hybrid · AI-start + tung redigering",
    "ren_ai": "Ren AI",
    "melding": "Uformelt · meldinger/notater",
}

HUMAN_CATEGORIES = {"ren_menneske_gammel", "ren_menneske_ny", "melding"}


class Sample(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    user_id: str
    title: str
    content: str
    source: Literal["paste", "file", "handwriting", "audio"] = "paste"
    category: str = "ren_menneske_ny"
    filename: Optional[str] = None
    word_count: int = 0
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))


class VoiceProfile(BaseModel):
    model_config = ConfigDict(extra="ignore")
    user_id: str
    updated_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    total_samples: int = 0
    total_words: int = 0
    avg_sentence_length: float = 0
    avg_word_length: float = 0
    vocabulary_richness: float = 0  # unique / total
    top_words: List[dict] = []
    function_word_frequencies: List[dict] = []
    sentence_length_distribution: List[dict] = []
    tone_description: str = ""
    style_summary: str = ""
    signature_phrases: List[str] = []
    # Runde 2 — stemmeprofilen som intellekt
    deviations: List[str] = []          # «Dette avviker» — mønstre som glir bort fra din vanlige stemme
    watch_out_for: List[str] = []       # «Dette bør du passe på» — konkrete ting å være oppmerksom på


ALLOWED_MODELS = {
    # Claude — Anthropic. claude-sonnet-4-6 er anbefalt for kreativ skriving.
    "claude-sonnet-4-5": ("anthropic", "claude-sonnet-4-5-20250929"),
    "claude-sonnet-4-6": ("anthropic", "claude-sonnet-4-6"),
    "claude-sonnet-5":   ("anthropic", "claude-sonnet-5"),
    "claude-opus-4-7":   ("anthropic", "claude-opus-4-7"),  # tunge analytiske oppgaver
    "claude-opus-5":     ("anthropic", "claude-opus-5"),
    # OpenAI
    "gpt-5.2": ("openai", "gpt-5.2"),
    "gpt-5.4": ("openai", "gpt-5.4"),
    # Google Gemini
    "gemini-3-pro":   ("gemini", "gemini-3.1-pro-preview"),
    "gemini-3-flash": ("gemini", "gemini-3-flash-preview"),
}

# Intern «beste-modell»-router. Bragarmål velger automatisk — brukeren ser ikke dette.
# Claude Sonnet 4.6 er anbefalt for kreativ norsk skriving; egne kanaler for spesifikke behov.
BRAGR_MODEL_WRITING = ("anthropic", "claude-sonnet-4-6")    # /generate + humaniser
BRAGR_MODEL_ANALYSIS = ("anthropic", "claude-sonnet-4-6")   # stemmeprofil, karakterer, AI-detektor
BRAGR_MODEL_VISION = ("anthropic", "claude-sonnet-4-6")     # OCR / håndskriftscan

# Direkte leverandørnøkler for /generate sin modell-nedtrekksliste (og BYOK-hjelpere
# lagrer sin egen nøkkel per provider, se helper["api_key"] i /generate).
PROVIDER_API_KEYS = {
    "anthropic": ANTHROPIC_API_KEY,
    "openai": OPENAI_API_KEY,
    "gemini": GEMINI_API_KEY,
    "xai": XAI_API_KEY,
}


async def stream_llm_text(
    provider: str, model: str, system: str, user_msg: str, api_key: str,
    temperature: float = 0.7, max_tokens: int = 4096, image_b64: Optional[str] = None,
):
    """Streams text deltas from any litellm-supported provider (anthropic, openai,
    gemini, xai, ...) using a caller-supplied API key. Callers are responsible for
    checking `api_key` is set before calling — this raises naturally otherwise."""
    content = user_msg
    if image_b64:
        content = [
            {"type": "text", "text": user_msg},
            {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{image_b64}"}},
        ]
    resp = await litellm.acompletion(
        model=f"{provider}/{model}",
        api_key=api_key,
        temperature=temperature,
        max_tokens=max_tokens,
        stream=True,
        messages=[
            {"role": "system", "content": system},
            {"role": "user", "content": content},
        ],
    )
    async for chunk in resp:
        delta = chunk.choices[0].delta.content if chunk.choices else None
        if delta:
            yield delta


# ---------- Auth ----------
async def get_current_user(request: Request) -> User:
    token = request.cookies.get("session_token")
    if not token:
        auth = request.headers.get("Authorization")
        if auth and auth.startswith("Bearer "):
            token = auth.split(" ", 1)[1]
    if not token:
        raise HTTPException(status_code=401, detail="Ikke autentisert")

    session = await db.user_sessions.find_one({"session_token": token}, {"_id": 0})
    if not session:
        raise HTTPException(status_code=401, detail="Ugyldig økt")

    expires_at = session.get("expires_at")
    if isinstance(expires_at, str):
        expires_at = datetime.fromisoformat(expires_at)
    if expires_at.tzinfo is None:
        expires_at = expires_at.replace(tzinfo=timezone.utc)
    if expires_at < datetime.now(timezone.utc):
        raise HTTPException(status_code=401, detail="Økten er utløpt")

    user_doc = await db.users.find_one({"user_id": session["user_id"]}, {"_id": 0})
    if not user_doc:
        raise HTTPException(status_code=401, detail="Bruker ikke funnet")
    return User(**user_doc)


def _hash_password(password: str) -> str:
    return bcrypt.hashpw(password.encode("utf-8"), bcrypt.gensalt()).decode("utf-8")


def _verify_password(password: str, hashed: str) -> bool:
    try:
        return bcrypt.checkpw(password.encode("utf-8"), hashed.encode("utf-8"))
    except Exception:
        return False


async def _create_session_for_user(user_id: str, response: Response) -> str:
    session_token = secrets.token_urlsafe(32)
    expires_at = datetime.now(timezone.utc) + timedelta(days=7)
    await db.user_sessions.insert_one({
        "user_id": user_id,
        "session_token": session_token,
        "expires_at": expires_at.isoformat(),
        "created_at": datetime.now(timezone.utc).isoformat(),
    })
    response.set_cookie(
        key="session_token",
        value=session_token,
        httponly=True,
        secure=True,
        samesite="none",
        max_age=7 * 24 * 60 * 60,
        path="/",
    )
    return session_token


class EmailRegister(BaseModel):
    email: str
    password: str
    name: Optional[str] = None


class EmailLogin(BaseModel):
    email: str
    password: str


@api_router.post("/auth/register")
async def auth_register(payload: EmailRegister, request: Request, response: Response):
    email = payload.email.strip().lower()
    if "@" not in email or len(email) < 5:
        raise HTTPException(status_code=400, detail="Ugyldig e-post")
    if len(payload.password) < 8:
        raise HTTPException(status_code=400, detail="Passord må være minst 8 tegn")

    existing = await db.users.find_one({"email": email}, {"_id": 0})
    if existing:
        raise HTTPException(status_code=400, detail="En bruker med denne e-posten finnes allerede")

    user_id = f"user_{uuid.uuid4().hex[:12]}"
    await db.users.insert_one({
        "user_id": user_id,
        "email": email,
        "name": (payload.name or email.split("@")[0]).strip()[:80],
        "picture": None,
        "password_hash": _hash_password(payload.password),
        "auth_provider": "email",
        "created_at": datetime.now(timezone.utc).isoformat(),
    })

    session_token = await _create_session_for_user(user_id, response)

    return {
        "user_id": user_id,
        "email": email,
        "name": (payload.name or email.split("@")[0]).strip()[:80],
        "picture": None,
        "session_token": session_token,
    }


@api_router.post("/auth/login")
async def auth_login(payload: EmailLogin, request: Request, response: Response):
    email = payload.email.strip().lower()
    ip = request.client.host if request.client else "unknown"
    identifier = f"{ip}:{email}"

    # Brute-force lockout: 5 failed attempts within 15 minutes
    now = datetime.now(timezone.utc)
    attempts_doc = await db.login_attempts.find_one({"identifier": identifier}, {"_id": 0})
    if attempts_doc and attempts_doc.get("locked_until"):
        try:
            locked_until = datetime.fromisoformat(attempts_doc["locked_until"])
            if locked_until.tzinfo is None:
                locked_until = locked_until.replace(tzinfo=timezone.utc)
            if locked_until > now:
                raise HTTPException(status_code=429, detail="For mange feilede forsøk. Prøv igjen om 15 minutter.")
        except HTTPException:
            raise
        except Exception:
            pass

    user_doc = await db.users.find_one({"email": email}, {"_id": 0})
    if not user_doc or not user_doc.get("password_hash"):
        # Increment failed attempts
        await _record_failed_login(identifier)
        raise HTTPException(status_code=401, detail="Ugyldig e-post eller passord")

    if not _verify_password(payload.password, user_doc["password_hash"]):
        await _record_failed_login(identifier)
        raise HTTPException(status_code=401, detail="Ugyldig e-post eller passord")

    # Success — clear attempts and create session
    await db.login_attempts.delete_one({"identifier": identifier})
    session_token = await _create_session_for_user(user_doc["user_id"], response)

    return {
        "user_id": user_doc["user_id"],
        "email": user_doc["email"],
        "name": user_doc.get("name"),
        "picture": user_doc.get("picture"),
        "session_token": session_token,
    }


async def _record_failed_login(identifier: str):
    now = datetime.now(timezone.utc)
    doc = await db.login_attempts.find_one({"identifier": identifier}, {"_id": 0})
    attempts = (doc or {}).get("attempts", 0) + 1
    update = {"identifier": identifier, "attempts": attempts, "last_attempt": now.isoformat()}
    if attempts >= 5:
        update["locked_until"] = (now + timedelta(minutes=15)).isoformat()
        update["attempts"] = 0  # reset after lock
    await db.login_attempts.update_one(
        {"identifier": identifier}, {"$set": update}, upsert=True
    )


async def _upsert_user_and_start_session(response: Response, email: str, name: str, picture: Optional[str]) -> dict:
    """Shared by every login path (Emergent-relayed or direct Google OAuth):
    upsert the user, mint a session token, set the cookie."""
    existing = await db.users.find_one({"email": email}, {"_id": 0})
    if existing:
        user_id = existing["user_id"]
        await db.users.update_one(
            {"user_id": user_id},
            {"$set": {"name": name, "picture": picture}},
        )
    else:
        user_id = f"user_{uuid.uuid4().hex[:12]}"
        await db.users.insert_one({
            "user_id": user_id,
            "email": email,
            "name": name,
            "picture": picture,
            "created_at": datetime.now(timezone.utc).isoformat(),
        })

    session_token = secrets.token_urlsafe(32)
    expires_at = datetime.now(timezone.utc) + timedelta(days=7)
    await db.user_sessions.insert_one({
        "user_id": user_id,
        "session_token": session_token,
        "expires_at": expires_at.isoformat(),
        "created_at": datetime.now(timezone.utc).isoformat(),
    })

    response.set_cookie(
        key="session_token",
        value=session_token,
        httponly=True,
        secure=True,
        samesite="none",
        max_age=7 * 24 * 60 * 60,
        path="/",
    )

    return {
        "user_id": user_id,
        "email": email,
        "name": name,
        "picture": picture,
        "session_token": session_token,
    }


@api_router.post("/auth/session")
async def create_session(request: Request, response: Response):
    """Exchange session_id from Emergent Auth for a session cookie."""
    body = await request.json()
    session_id = body.get("session_id")
    if not session_id:
        raise HTTPException(status_code=400, detail="Mangler session_id")

    async with httpx.AsyncClient(timeout=15) as h:
        r = await h.get(
            "https://demobackend.emergentagent.com/auth/v1/env/oauth/session-data",
            headers={"X-Session-ID": session_id},
        )
        if r.status_code != 200:
            raise HTTPException(status_code=401, detail="Kunne ikke bekrefte økt")
        data = r.json()

    result = await _upsert_user_and_start_session(
        response, data["email"], data.get("name", data["email"]), data.get("picture")
    )
    result.pop("session_token", None)
    return result


# ---------- Direct Google OAuth (Emergent-independent login) ----------
# Only active once GOOGLE_OAUTH_CLIENT_ID/SECRET are configured — otherwise these
# routes 404 and the Emergent-relayed flow above keeps working unchanged.
GOOGLE_OAUTH_CLIENT_ID = os.environ.get('GOOGLE_OAUTH_CLIENT_ID')
GOOGLE_OAUTH_CLIENT_SECRET = os.environ.get('GOOGLE_OAUTH_CLIENT_SECRET')
GOOGLE_OAUTH_REDIRECT_URI = os.environ.get('GOOGLE_OAUTH_REDIRECT_URI')  # e.g. https://api.bragarmål.no/api/auth/google/callback
FRONTEND_URL = os.environ.get('FRONTEND_URL', 'https://xn--bragarml-g0a.no')

_oauth_states: dict = {}  # state -> created_at (short-lived, in-memory CSRF guard)


@api_router.get("/auth/google/start")
async def google_login_start():
    if not (GOOGLE_OAUTH_CLIENT_ID and GOOGLE_OAUTH_REDIRECT_URI):
        raise HTTPException(status_code=503, detail="Direkte Google-innlogging er ikke konfigurert ennå")
    state = secrets.token_urlsafe(24)
    _oauth_states[state] = datetime.now(timezone.utc)
    # Prune anything older than 10 minutes so this dict never grows unbounded
    cutoff = datetime.now(timezone.utc) - timedelta(minutes=10)
    for k in [k for k, v in _oauth_states.items() if v < cutoff]:
        _oauth_states.pop(k, None)

    params = {
        "client_id": GOOGLE_OAUTH_CLIENT_ID,
        "redirect_uri": GOOGLE_OAUTH_REDIRECT_URI,
        "response_type": "code",
        "scope": "openid email profile",
        "state": state,
        "access_type": "online",
        "prompt": "select_account",
    }
    return JSONResponse({"auth_url": f"https://accounts.google.com/o/oauth2/v2/auth?{urlencode(params)}"})


@api_router.get("/auth/google/callback")
async def google_login_callback(response: Response, code: Optional[str] = None, state: Optional[str] = None, error: Optional[str] = None):
    if error or not code or not state or state not in _oauth_states:
        return RedirectResponse(f"{FRONTEND_URL}/logg-inn?feil=innlogging_avbrutt")
    _oauth_states.pop(state, None)

    async with httpx.AsyncClient(timeout=15) as h:
        token_r = await h.post("https://oauth2.googleapis.com/token", data={
            "code": code,
            "client_id": GOOGLE_OAUTH_CLIENT_ID,
            "client_secret": GOOGLE_OAUTH_CLIENT_SECRET,
            "redirect_uri": GOOGLE_OAUTH_REDIRECT_URI,
            "grant_type": "authorization_code",
        })
        if token_r.status_code != 200:
            logger.error("Google token exchange failed: status=%s body=%s redirect_uri=%s",
                         token_r.status_code, token_r.text, GOOGLE_OAUTH_REDIRECT_URI)
            return RedirectResponse(f"{FRONTEND_URL}/logg-inn?feil=google_token")
        access_token = token_r.json()["access_token"]

        userinfo_r = await h.get(
            "https://www.googleapis.com/oauth2/v3/userinfo",
            headers={"Authorization": f"Bearer {access_token}"},
        )
        if userinfo_r.status_code != 200:
            logger.error("Google userinfo fetch failed: status=%s body=%s", userinfo_r.status_code, userinfo_r.text)
            return RedirectResponse(f"{FRONTEND_URL}/logg-inn?feil=google_userinfo")
        profile = userinfo_r.json()

    result = await _upsert_user_and_start_session(
        response, profile["email"], profile.get("name", profile["email"]), profile.get("picture")
    )
    # Cookie is set above for browsers that accept it, but the frontend and
    # backend are on different sites, so third-party-cookie blocking can
    # silently drop it. Pass the token in the URL hash too (never sent to
    # any server, incl. ours) so the frontend can store it and send it back
    # as an Authorization header instead.
    redirect = RedirectResponse(f"{FRONTEND_URL}/dashboard#session_token={result['session_token']}")
    for k, v in response.headers.items():
        if k.lower() == "set-cookie":
            redirect.headers.append(k, v)
    # The cookie must be set directly on the object we return — FastAPI discards
    # the injected `response` param's headers once a Response is returned explicitly.
    return redirect


@api_router.get("/auth/me")
async def me(user: User = Depends(get_current_user)):
    return {
        "user_id": user.user_id,
        "email": user.email,
        "name": user.name,
        "picture": user.picture,
    }


@api_router.post("/auth/logout")
async def logout(request: Request, response: Response):
    token = request.cookies.get("session_token")
    if token:
        await db.user_sessions.delete_one({"session_token": token})
    response.delete_cookie("session_token", path="/")
    return {"ok": True}


# ---------- File extraction ----------
def extract_text_from_upload(filename: str, data: bytes) -> str:
    name = filename.lower()
    if name.endswith(".txt") or name.endswith(".md"):
        try:
            return data.decode("utf-8")
        except UnicodeDecodeError:
            return data.decode("latin-1", errors="ignore")
    if name.endswith(".pdf"):
        reader = PdfReader(io.BytesIO(data))
        return "\n\n".join((p.extract_text() or "") for p in reader.pages)
    if name.endswith(".docx"):
        doc = DocxDocument(io.BytesIO(data))
        return "\n\n".join(p.text for p in doc.paragraphs)
    raise HTTPException(status_code=400, detail="Filtype støttes ikke. Bruk .txt, .pdf eller .docx")


# ---------- Voice analysis ----------
NORWEGIAN_STOPWORDS = {
    "og","i","jeg","det","at","en","et","den","til","er","som","på","de","med","han",
    "av","ikke","der","så","var","meg","seg","men","ett","har","om","vi","min","mitt",
    "ha","hadde","hun","nå","over","da","ved","fra","du","ut","sin","dem","oss","opp",
    "man","kan","hans","hvor","eller","hva","skal","selv","sjøl","her","alle","vil",
    "bli","ble","blitt","kunne","inn","når","være","kom","noen","noe","ville","dere",
    "som","deres","kun","ja","etter","ned","skulle","denne","for","deg","si","sine",
    "sitt","mot","å","meget","hvorfor","dette","disse","uten","hvordan","ingen","din",
    "ditt","blir","samme","hvilken","hvilke","sånn","inni","mellom","vår","hver",
    "hvem","vors","hvis","både","bare","enn","fordi","før","mange","også","slik",
    "vært","være","båe","begge","siden","dykk","dykkar","dei","deira","deires","egen",
    "et","hu","enno","annet","annen","andre","the"
}


def analyze_voice(samples: List[dict]) -> dict:
    """Local, deterministic analysis of writing style."""
    text = "\n\n".join(s["content"] for s in samples)
    words_all = re.findall(r"[A-Za-zÆØÅæøåéèê']+", text)
    if not words_all:
        return {
            "total_words": 0,
            "avg_sentence_length": 0,
            "avg_word_length": 0,
            "vocabulary_richness": 0,
            "top_words": [],
            "sentence_length_distribution": [],
        }

    words_lower = [w.lower() for w in words_all]
    unique = set(words_lower)
    avg_word_len = sum(len(w) for w in words_all) / len(words_all)
    richness = len(unique) / len(words_lower)

    # Sentences
    sentences = re.split(r"(?<=[.!?])\s+", text.strip())
    sentences = [s for s in sentences if s.strip()]
    sent_lens = [len(re.findall(r"[A-Za-zÆØÅæøå']+", s)) for s in sentences]
    sent_lens = [n for n in sent_lens if n > 0]
    avg_sent = sum(sent_lens) / len(sent_lens) if sent_lens else 0

    # Distribution buckets
    buckets = [
        ("1-5", 0), ("6-10", 0), ("11-15", 0), ("16-20", 0), ("21-30", 0), ("31+", 0)
    ]
    b = dict(buckets)
    for n in sent_lens:
        if n <= 5: b["1-5"] += 1
        elif n <= 10: b["6-10"] += 1
        elif n <= 15: b["11-15"] += 1
        elif n <= 20: b["16-20"] += 1
        elif n <= 30: b["21-30"] += 1
        else: b["31+"] += 1
    distribution = [{"range": k, "count": v} for k, v in b.items()]

    # Top words (excluding stopwords, length >= 4)
    freq = {}
    for w in words_lower:
        if w in NORWEGIAN_STOPWORDS or len(w) < 4:
            continue
        freq[w] = freq.get(w, 0) + 1
    top = sorted(freq.items(), key=lambda x: -x[1])[:15]
    top_words = [{"word": w, "count": c} for w, c in top]

    # Function-word frequencies (per 1000 words) — strong stylometric marker
    fw_targets = [
        "og","men","som","at","fordi","hvis","når","der","hvor","der",
        "ikke","også","kanskje","liksom","altså","bare","enda","jo","nok","alltid",
        "kanskje","selv","ganske","helt","litt","veldig","så","da","mens","selv",
    ]
    fw_set = list(dict.fromkeys(fw_targets))
    fw_counts = {}
    for w in words_lower:
        if w in fw_set:
            fw_counts[w] = fw_counts.get(w, 0) + 1
    per_1000 = 1000.0 / max(len(words_lower), 1)
    function_word_frequencies = sorted(
        [{"word": w, "per_1000": round(c * per_1000, 2), "count": c} for w, c in fw_counts.items()],
        key=lambda x: -x["per_1000"],
    )[:20]

    return {
        "total_words": len(words_lower),
        "avg_sentence_length": round(avg_sent, 2),
        "avg_word_length": round(avg_word_len, 2),
        "vocabulary_richness": round(richness, 3),
        "top_words": top_words,
        "function_word_frequencies": function_word_frequencies,
        "sentence_length_distribution": distribution,
    }


async def build_style_summary(samples: List[dict], stats: dict) -> dict:
    """Use Claude to describe user's tone and signature phrases in Norwegian."""
    if not samples or not ANTHROPIC_API_KEY:
        return {"tone_description": "", "style_summary": "", "signature_phrases": [], "deviations": [], "watch_out_for": []}

    # Ensure ALL samples are represented — take proportional slices from each
    # (avoids the bug where newest samples get truncated away by a naive [:12000] cut)
    MAX_TOTAL = 12000
    if samples:
        per_sample = max(200, MAX_TOTAL // len(samples))
        pieces = []
        for s in samples:
            c = (s.get("content") or "").strip()
            if len(c) > per_sample:
                c = c[:per_sample].rsplit(" ", 1)[0] + "…"
            pieces.append(c)
        joined = "\n\n---\n\n".join(pieces)
        if len(joined) > MAX_TOTAL:
            joined = joined[:MAX_TOTAL]
    else:
        joined = ""

    system = (
        "Du er en litterær stilanalytiker. Du analyserer en forfatters norske tekster "
        "og beskriver stemmen kort, konkret og uten klisjeer. Svar KUN i gyldig JSON."
    )
    prompt = f"""Analyser følgende tekstprøver fra samme forfatter. Basert på faktiske detaljer i teksten:

1. `tone_description`: 1-2 setninger på norsk som beskriver tonen (rytme, temperatur, distanse, humor, alvor).
2. `style_summary`: 2-3 setninger på norsk om setningsbygning, ordvalg og typiske grep.
3. `signature_phrases`: 5-8 karakteristiske uttrykk, ord eller vendinger forfatteren faktisk bruker (kopier direkte fra tekstene).
4. `deviations`: 3-5 korte observasjoner om hvor teksten noen ganger glir vekk fra forfatterens vanlige stemme. Konkret. F.eks. «Passasjer om arbeid glir mot mer formell tone», «Følelsesmessig ladede scener bruker lengre setninger enn hennes vanlige knappe rytme», «Enkelte dialoger virker skrevet for høyt — mister det hverdagslige som ellers preger stemmen». Ikke råd — bare observasjoner.
5. `watch_out_for`: 3-5 konkrete ting forfatteren kan være oppmerksom på i eget arbeid. Ikke pekefinger — vennlige merknader. F.eks. «Gjentakelser av ordet 'liksom' i påfølgende setninger», «Tendens til å forklare følelser i stedet for å vise dem», «Komma-feilstillinger i lange innskudd», «Overgangsfraser som 'og så' som lener rytmen».

Statistikk: gjennomsnittlig setningslengde {stats.get('avg_sentence_length')} ord, ordforråd-rikhet {stats.get('vocabulary_richness')}.

Svar i dette JSON-formatet, ingen forklaring utenfor JSON:
{{"tone_description": "...", "style_summary": "...", "signature_phrases": ["...", "..."], "deviations": ["...", "..."], "watch_out_for": ["...", "..."]}}

TEKSTER:
{joined}
"""

    try:
        parts = []
        async for chunk in stream_llm_text(*BRAGR_MODEL_ANALYSIS, system, prompt, ANTHROPIC_API_KEY):
            parts.append(chunk)
        raw = "".join(parts).strip()
        # Extract JSON block
        m = re.search(r"\{.*\}", raw, re.DOTALL)
        if m:
            raw = m.group(0)
        parsed = json.loads(raw)
        return {
            "tone_description": parsed.get("tone_description", ""),
            "style_summary": parsed.get("style_summary", ""),
            "signature_phrases": parsed.get("signature_phrases", [])[:8],
            "deviations": parsed.get("deviations", [])[:6],
            "watch_out_for": parsed.get("watch_out_for", [])[:6],
        }
    except Exception as e:
        logger.warning(f"Style summary failed: {e}")
        return {"tone_description": "", "style_summary": "", "signature_phrases": [], "deviations": [], "watch_out_for": []}


# ---------- Sample endpoints ----------
class SampleCreate(BaseModel):
    title: str
    content: str
    category: Optional[str] = None


@api_router.post("/samples")
async def create_sample(payload: SampleCreate, user: User = Depends(get_current_user)):
    content = payload.content.strip()
    if len(content) < 20:
        raise HTTPException(status_code=400, detail="Teksten er for kort (minst 20 tegn)")
    category = payload.category if payload.category in SAMPLE_CATEGORIES else "ren_menneske_ny"
    sample = Sample(
        user_id=user.user_id,
        title=payload.title.strip() or "Uten tittel",
        content=content,
        source="paste",
        category=category,
        word_count=len(re.findall(r"\S+", content)),
    )
    doc = sample.model_dump()
    doc["created_at"] = doc["created_at"].isoformat()
    await db.samples.insert_one(doc)
    return sample


@api_router.post("/samples/upload")
async def upload_sample(
    file: UploadFile = File(...),
    title: Optional[str] = Form(None),
    category: Optional[str] = Form(None),
    user: User = Depends(get_current_user),
):
    data = await file.read()
    if len(data) > 5 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="Filen er for stor (maks 5 MB)")
    text = extract_text_from_upload(file.filename, data)
    text = text.strip()
    if len(text) < 20:
        raise HTTPException(status_code=400, detail="Kunne ikke lese meningsfull tekst fra filen")
    cat = category if category in SAMPLE_CATEGORIES else "ren_menneske_ny"

    # Store original file in object storage
    storage_path = None
    try:
        ext = file.filename.rsplit(".", 1)[-1].lower() if "." in file.filename else "bin"
        path = f"{APP_NAME}/uploads/{user.user_id}/{uuid.uuid4()}.{ext}"
        result = storage_put(path, data, file.content_type or "application/octet-stream")
        storage_path = result["path"]
        file_id = str(uuid.uuid4())
        await db.files.insert_one({
            "id": file_id,
            "user_id": user.user_id,
            "storage_path": storage_path,
            "original_filename": file.filename,
            "content_type": file.content_type or "application/octet-stream",
            "size": len(data),
            "kind": "document",
            "is_deleted": False,
            "created_at": datetime.now(timezone.utc).isoformat(),
        })
    except Exception as e:
        logger.warning(f"Original file not stored: {e}")
        file_id = None

    sample = Sample(
        user_id=user.user_id,
        title=(title or file.filename or "Uten tittel").strip(),
        content=text,
        source="file",
        category=cat,
        filename=file.filename,
        word_count=len(re.findall(r"\S+", text)),
    )
    doc = sample.model_dump()
    doc["created_at"] = doc["created_at"].isoformat()
    if file_id:
        doc["file_id"] = file_id
    await db.samples.insert_one(doc)
    return sample


def suggest_title(text: str) -> str:
    """First 6-8 meaningful words of the transcript as a title."""
    if not text:
        return ""
    first_line = text.strip().split("\n", 1)[0].strip()
    if len(first_line) < 8:
        # merge with next line
        parts = text.strip().split("\n", 2)
        first_line = " ".join(p.strip() for p in parts[:2] if p.strip())
    words = re.findall(r"\S+", first_line)
    title = " ".join(words[:8])
    if len(title) > 60:
        title = title[:57].rsplit(" ", 1)[0] + "…"
    return title or "Uten tittel"


@api_router.post("/samples/scan")
async def scan_handwritten(
    file: UploadFile = File(...),
    user: User = Depends(get_current_user),
):
    """OCR a photo of handwritten Norwegian text using Claude Sonnet 4.5 vision.
    Returns extracted text without saving as a sample — user can review before saving.
    """
    import base64

    if not ANTHROPIC_API_KEY:
        raise HTTPException(status_code=500, detail="LLM-nøkkel mangler")

    ct = (file.content_type or "").lower()
    if not (ct.startswith("image/") or file.filename.lower().endswith((".jpg", ".jpeg", ".png", ".webp", ".heic"))):
        raise HTTPException(status_code=400, detail="Bare bildefiler (.jpg, .png, .webp)")

    data = await file.read()
    if len(data) > 15 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="Bildet er for stort (maks 15 MB)")

    b64 = base64.b64encode(data).decode("utf-8")

    system = (
        "Du er en nøyaktig transkriberer av norsk håndskrift. Din oppgave er å transkribere "
        "teksten i bildet HELT ORDRETT — bevar setningsstruktur, linjeskift, og alle særegenheter "
        "i skrivemåten (staveformer, dialektord, gammeldags skrivemåte). "
        "IKKE forbedre grammatikken. IKKE moderniser språket. IKKE oppsummer. "
        "Hvis noen ord er uleselige, marker med [uleselig] i stedet for å gjette. "
        "Returner KUN den transkriberte teksten, ingen forklaring."
    )

    try:
        parts = []
        async for chunk in stream_llm_text(
            *BRAGR_MODEL_VISION, system,
            "Transkriber all håndskrevet norsk tekst i dette bildet ordrett.",
            ANTHROPIC_API_KEY, image_b64=b64,
        ):
            parts.append(chunk)
        text = "".join(parts).strip()
    except Exception as e:
        logger.exception("OCR failed")
        raise HTTPException(status_code=500, detail=f"Kunne ikke transkribere: {e}")

    if len(text) < 5:
        raise HTTPException(status_code=400, detail="Fant ikke lesbar tekst i bildet")

    # Store original image
    file_id = None
    try:
        ext = (file.filename or "photo.jpg").rsplit(".", 1)[-1].lower() if "." in (file.filename or "") else "jpg"
        path = f"{APP_NAME}/uploads/{user.user_id}/{uuid.uuid4()}.{ext}"
        result = storage_put(path, data, file.content_type or "image/jpeg")
        file_id = str(uuid.uuid4())
        await db.files.insert_one({
            "id": file_id,
            "user_id": user.user_id,
            "storage_path": result["path"],
            "original_filename": file.filename or "handwriting.jpg",
            "content_type": file.content_type or "image/jpeg",
            "size": len(data),
            "kind": "handwriting",
            "is_deleted": False,
            "created_at": datetime.now(timezone.utc).isoformat(),
        })
    except Exception as e:
        logger.warning(f"Handwriting image not stored: {e}")

    return {
        "text": text,
        "word_count": len(re.findall(r"\S+", text)),
        "filename": file.filename,
        "suggested_title": suggest_title(text),
        "file_id": file_id,
    }


@api_router.post("/samples/transcribe")
async def transcribe_audio(
    file: UploadFile = File(...),
    user: User = Depends(get_current_user),
):
    """Transcribe a recorded Norwegian audio file (høytlesning) to text via Whisper.
    Returns the transcript — user reviews and can save as a sample.
    """
    if not OPENAI_API_KEY:
        raise HTTPException(status_code=500, detail="LLM-nøkkel mangler")

    name = (file.filename or "").lower()
    ct = (file.content_type or "").lower()
    if not (ct.startswith("audio/") or ct.startswith("video/webm") or
            name.endswith((".mp3", ".mp4", ".mpeg", ".mpga", ".m4a", ".wav", ".webm"))):
        raise HTTPException(status_code=400, detail="Bare lydfiler (.mp3, .m4a, .wav, .webm)")

    data = await file.read()
    if len(data) > 25 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="Lyden er for stor (maks 25 MB)")

    # Write to a temp file so OpenAI SDK can read it
    import tempfile
    suffix = os.path.splitext(name)[1] or ".webm"
    with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
        tmp.write(data)
        tmp_path = tmp.name

    try:
        client = AsyncOpenAI(api_key=OPENAI_API_KEY)
        with open(tmp_path, "rb") as af:
            resp = await client.audio.transcriptions.create(
                file=af,
                model="whisper-1",
                response_format="json",
                language="no",
                temperature=0.0,
                prompt="Norsk høytlesning fra forfatter. Bevar egne ord og dialekt.",
            )
        text = (getattr(resp, "text", None) or "").strip()
    except Exception as e:
        logger.exception("Whisper failed")
        raise HTTPException(status_code=500, detail=f"Kunne ikke transkribere lyd: {e}")
    finally:
        try: os.unlink(tmp_path)
        except Exception: pass

    if len(text) < 3:
        raise HTTPException(status_code=400, detail="Fant ingen tale i lyden")

    # Store original audio
    file_id = None
    try:
        ext = (file.filename or "audio.webm").rsplit(".", 1)[-1].lower() if "." in (file.filename or "") else "webm"
        path = f"{APP_NAME}/uploads/{user.user_id}/{uuid.uuid4()}.{ext}"
        result = storage_put(path, data, file.content_type or "audio/webm")
        file_id = str(uuid.uuid4())
        await db.files.insert_one({
            "id": file_id,
            "user_id": user.user_id,
            "storage_path": result["path"],
            "original_filename": file.filename or "audio.webm",
            "content_type": file.content_type or "audio/webm",
            "size": len(data),
            "kind": "audio",
            "is_deleted": False,
            "created_at": datetime.now(timezone.utc).isoformat(),
        })
    except Exception as e:
        logger.warning(f"Audio not stored: {e}")

    return {
        "text": text,
        "word_count": len(re.findall(r"\S+", text)),
        "filename": file.filename,
        "suggested_title": suggest_title(text),
        "file_id": file_id,
    }


@api_router.get("/samples")
async def list_samples(user: User = Depends(get_current_user)):
    docs = await db.samples.find({"user_id": user.user_id}, {"_id": 0}).sort("created_at", -1).to_list(200)
    for d in docs:
        if isinstance(d.get("created_at"), str):
            pass
    return docs


@api_router.delete("/samples/{sample_id}")
async def delete_sample(sample_id: str, user: User = Depends(get_current_user)):
    r = await db.samples.delete_one({"id": sample_id, "user_id": user.user_id})
    if r.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Ikke funnet")
    return {"ok": True}


class SampleUpdate(BaseModel):
    title: Optional[str] = None
    content: Optional[str] = None


@api_router.patch("/samples/{sample_id}")
async def update_sample(
    sample_id: str,
    body: SampleUpdate,
    user: User = Depends(get_current_user),
):
    updates: dict = {}
    if body.title is not None:
        updates["title"] = body.title.strip() or "Uten tittel"
    if body.content is not None:
        content = body.content.strip()
        if len(content) < 20:
            raise HTTPException(status_code=400, detail="Innholdet må være minst 20 tegn")
        updates["content"] = content
        updates["word_count"] = len(content.split())
    if not updates:
        raise HTTPException(status_code=400, detail="Ingenting å oppdatere")
    r = await db.samples.update_one(
        {"id": sample_id, "user_id": user.user_id},
        {"$set": updates},
    )
    if r.matched_count == 0:
        raise HTTPException(status_code=404, detail="Ikke funnet")
    doc = await db.samples.find_one(
        {"id": sample_id, "user_id": user.user_id}, {"_id": 0}
    )
    return doc


class SampleCategoryUpdate(BaseModel):
    category: str


@api_router.patch("/samples/{sample_id}/category")
async def update_sample_category(
    sample_id: str,
    body: SampleCategoryUpdate,
    user: User = Depends(get_current_user),
):
    if body.category not in SAMPLE_CATEGORIES:
        raise HTTPException(status_code=400, detail="Ukjent kategori")
    r = await db.samples.update_one(
        {"id": sample_id, "user_id": user.user_id},
        {"$set": {"category": body.category}},
    )
    if r.matched_count == 0:
        raise HTTPException(status_code=404, detail="Ikke funnet")
    return {"ok": True, "category": body.category}


@api_router.get("/categories")
async def list_categories():
    return [{"id": k, "label": v} for k, v in SAMPLE_CATEGORIES.items()]


# ---------- Manuscript / Outliner (scenes & chapters) ----------

SCENE_STATUSES = ["skisse", "utkast", "revidert", "ferdig"]


class SceneCreate(BaseModel):
    title: str = "Uten tittel"
    synopsis: str = ""
    content: str = ""
    status: str = "skisse"
    pov: str = ""
    location: str = ""
    scene_date: str = ""  # freeform user-entered date/timeline label
    tags: List[str] = []


class SceneUpdate(BaseModel):
    title: Optional[str] = None
    synopsis: Optional[str] = None
    content: Optional[str] = None
    status: Optional[str] = None
    pov: Optional[str] = None
    location: Optional[str] = None
    scene_date: Optional[str] = None
    tags: Optional[List[str]] = None
    scrapped: Optional[bool] = None


class SceneReorder(BaseModel):
    ordered_ids: List[str]


def _word_count(text: str) -> int:
    return len(text.strip().split()) if text and text.strip() else 0


@api_router.get("/manuscript")
async def list_scenes(user: User = Depends(get_current_user)):
    docs = await db.scenes.find({"user_id": user.user_id}, {"_id": 0}).sort("order", 1).to_list(1000)
    return docs


@api_router.post("/manuscript")
async def create_scene(body: SceneCreate, user: User = Depends(get_current_user)):
    if body.status not in SCENE_STATUSES:
        raise HTTPException(status_code=400, detail="Ugyldig status")
    # Find current max order
    last = await db.scenes.find({"user_id": user.user_id}, {"order": 1, "_id": 0}).sort("order", -1).limit(1).to_list(1)
    next_order = (last[0]["order"] + 1) if last else 0
    scene = {
        "id": str(uuid.uuid4()),
        "user_id": user.user_id,
        "title": body.title.strip() or "Uten tittel",
        "synopsis": body.synopsis.strip(),
        "content": body.content,
        "status": body.status,
        "pov": body.pov.strip(),
        "location": body.location.strip(),
        "scene_date": body.scene_date.strip(),
        "tags": [t.strip() for t in body.tags if t.strip()][:12],
        "scrapped": False,
        "word_count": _word_count(body.content),
        "order": next_order,
        "created_at": datetime.now(timezone.utc).isoformat(),
        "updated_at": datetime.now(timezone.utc).isoformat(),
    }
    await db.scenes.insert_one(scene)
    scene.pop("_id", None)
    return scene


@api_router.patch("/manuscript/{scene_id}")
async def update_scene(scene_id: str, body: SceneUpdate, user: User = Depends(get_current_user)):
    updates: dict = {}
    if body.title is not None:
        updates["title"] = body.title.strip() or "Uten tittel"
    if body.synopsis is not None:
        updates["synopsis"] = body.synopsis.strip()
    if body.content is not None:
        updates["content"] = body.content
        updates["word_count"] = _word_count(body.content)
    if body.status is not None:
        if body.status not in SCENE_STATUSES:
            raise HTTPException(status_code=400, detail="Ugyldig status")
        updates["status"] = body.status
    if body.pov is not None:
        updates["pov"] = body.pov.strip()
    if body.location is not None:
        updates["location"] = body.location.strip()
    if body.scene_date is not None:
        updates["scene_date"] = body.scene_date.strip()
    if body.tags is not None:
        updates["tags"] = [t.strip() for t in body.tags if t.strip()][:12]
    if body.scrapped is not None:
        updates["scrapped"] = body.scrapped
    if not updates:
        raise HTTPException(status_code=400, detail="Ingenting å oppdatere")
    updates["updated_at"] = datetime.now(timezone.utc).isoformat()
    r = await db.scenes.update_one(
        {"id": scene_id, "user_id": user.user_id},
        {"$set": updates},
    )
    if r.matched_count == 0:
        raise HTTPException(status_code=404, detail="Ikke funnet")
    doc = await db.scenes.find_one({"id": scene_id, "user_id": user.user_id}, {"_id": 0})
    return doc


@api_router.delete("/manuscript/{scene_id}")
async def delete_scene(scene_id: str, user: User = Depends(get_current_user)):
    r = await db.scenes.delete_one({"id": scene_id, "user_id": user.user_id})
    if r.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Ikke funnet")
    return {"ok": True}


@api_router.post("/manuscript/reorder")
async def reorder_scenes(body: SceneReorder, user: User = Depends(get_current_user)):
    # Bulk update all order values in a single pass
    for idx, scene_id in enumerate(body.ordered_ids):
        await db.scenes.update_one(
            {"id": scene_id, "user_id": user.user_id},
            {"$set": {"order": idx, "updated_at": datetime.now(timezone.utc).isoformat()}},
        )
    return {"ok": True, "count": len(body.ordered_ids)}


@api_router.get("/manuscript/statuses")
async def list_statuses():
    return SCENE_STATUSES


# ---------- Characters (psychological profiles) ----------
class CharacterCreate(BaseModel):
    name: str
    role: str = ""
    appearance: str = ""
    inner_struggle: str = ""
    outer_struggle: str = ""
    relationships: str = ""
    arc: str = ""
    voice_notes: str = ""


class CharacterUpdate(BaseModel):
    name: Optional[str] = None
    role: Optional[str] = None
    appearance: Optional[str] = None
    inner_struggle: Optional[str] = None
    outer_struggle: Optional[str] = None
    relationships: Optional[str] = None
    arc: Optional[str] = None
    voice_notes: Optional[str] = None


@api_router.get("/characters")
async def list_characters(user: User = Depends(get_current_user)):
    return await db.characters.find({"user_id": user.user_id}, {"_id": 0}).sort("name", 1).to_list(500)


@api_router.post("/characters")
async def create_character(body: CharacterCreate, user: User = Depends(get_current_user)):
    doc = {
        "id": str(uuid.uuid4()),
        "user_id": user.user_id,
        "name": body.name.strip() or "Uten navn",
        "role": body.role.strip(),
        "appearance": body.appearance.strip(),
        "inner_struggle": body.inner_struggle.strip(),
        "outer_struggle": body.outer_struggle.strip(),
        "relationships": body.relationships.strip(),
        "arc": body.arc.strip(),
        "voice_notes": body.voice_notes.strip(),
        "created_at": datetime.now(timezone.utc).isoformat(),
        "updated_at": datetime.now(timezone.utc).isoformat(),
    }
    await db.characters.insert_one(doc)
    doc.pop("_id", None)
    return doc


@api_router.patch("/characters/{char_id}")
async def update_character(char_id: str, body: CharacterUpdate, user: User = Depends(get_current_user)):
    updates = {k: (v.strip() if isinstance(v, str) else v) for k, v in body.model_dump(exclude_none=True).items()}
    if not updates:
        raise HTTPException(status_code=400, detail="Ingenting å oppdatere")
    updates["updated_at"] = datetime.now(timezone.utc).isoformat()
    r = await db.characters.update_one(
        {"id": char_id, "user_id": user.user_id}, {"$set": updates}
    )
    if r.matched_count == 0:
        raise HTTPException(status_code=404, detail="Ikke funnet")
    return await db.characters.find_one({"id": char_id, "user_id": user.user_id}, {"_id": 0})


@api_router.delete("/characters/{char_id}")
async def delete_character(char_id: str, user: User = Depends(get_current_user)):
    r = await db.characters.delete_one({"id": char_id, "user_id": user.user_id})
    if r.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Ikke funnet")
    return {"ok": True}


# ─── Faktastasjon (research notes) ──────────────────────────────────────────
RESEARCH_CATEGORIES = ["person", "sted", "tidsperiode", "gjenstand", "annet"]
RESEARCH_IMAGE_MAX_UPLOAD_BYTES = 8 * 1024 * 1024
RESEARCH_IMAGE_MAX_DIM = 1400


class ResearchNoteCreate(BaseModel):
    title: str
    category: str = "annet"
    content: str = ""
    source_url: str = ""


class ResearchNoteUpdate(BaseModel):
    title: Optional[str] = None
    category: Optional[str] = None
    content: Optional[str] = None
    source_url: Optional[str] = None


@api_router.get("/research")
async def list_research_notes(user: User = Depends(get_current_user)):
    return await db.research_notes.find(
        {"user_id": user.user_id}, {"_id": 0, "image_data": 0}
    ).sort("updated_at", -1).to_list(1000)


@api_router.post("/research")
async def create_research_note(body: ResearchNoteCreate, user: User = Depends(get_current_user)):
    category = body.category if body.category in RESEARCH_CATEGORIES else "annet"
    doc = {
        "id": str(uuid.uuid4()),
        "user_id": user.user_id,
        "title": body.title.strip() or "Uten tittel",
        "category": category,
        "content": body.content.strip(),
        "source_url": body.source_url.strip(),
        "has_image": False,
        "created_at": datetime.now(timezone.utc).isoformat(),
        "updated_at": datetime.now(timezone.utc).isoformat(),
    }
    await db.research_notes.insert_one(doc)
    doc.pop("_id", None)
    return doc


@api_router.patch("/research/{note_id}")
async def update_research_note(note_id: str, body: ResearchNoteUpdate, user: User = Depends(get_current_user)):
    updates = {k: (v.strip() if isinstance(v, str) else v) for k, v in body.model_dump(exclude_none=True).items()}
    if "category" in updates and updates["category"] not in RESEARCH_CATEGORIES:
        updates["category"] = "annet"
    if not updates:
        raise HTTPException(status_code=400, detail="Ingenting å oppdatere")
    updates["updated_at"] = datetime.now(timezone.utc).isoformat()
    r = await db.research_notes.update_one(
        {"id": note_id, "user_id": user.user_id}, {"$set": updates}
    )
    if r.matched_count == 0:
        raise HTTPException(status_code=404, detail="Ikke funnet")
    return await db.research_notes.find_one(
        {"id": note_id, "user_id": user.user_id}, {"_id": 0, "image_data": 0}
    )


@api_router.delete("/research/{note_id}")
async def delete_research_note(note_id: str, user: User = Depends(get_current_user)):
    r = await db.research_notes.delete_one({"id": note_id, "user_id": user.user_id})
    if r.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Ikke funnet")
    return {"ok": True}


@api_router.post("/research/{note_id}/image")
async def upload_research_note_image(note_id: str, file: UploadFile = File(...), user: User = Depends(get_current_user)):
    exists = await db.research_notes.find_one({"id": note_id, "user_id": user.user_id}, {"_id": 0, "id": 1})
    if not exists:
        raise HTTPException(status_code=404, detail="Notat ikke funnet")
    ct = (file.content_type or "").lower()
    if not ct.startswith("image/"):
        raise HTTPException(status_code=400, detail="Filen må være et bilde (PNG, JPG, WebP e.l.)")
    data = await file.read()
    if not data:
        raise HTTPException(status_code=400, detail="Tom fil")
    if len(data) > RESEARCH_IMAGE_MAX_UPLOAD_BYTES:
        raise HTTPException(status_code=400, detail="Bildet er for stort (maks 8 MB)")
    try:
        img = Image.open(io.BytesIO(data))
        img.load()
    except Exception:
        raise HTTPException(status_code=400, detail="Kunne ikke lese bildet — prøv en annen fil")
    if img.mode in ("RGBA", "LA", "P"):
        rgba = img.convert("RGBA")
        bg = Image.new("RGB", rgba.size, (255, 255, 255))
        bg.paste(rgba, mask=rgba.split()[-1])
        img = bg
    else:
        img = img.convert("RGB")
    img.thumbnail((RESEARCH_IMAGE_MAX_DIM, RESEARCH_IMAGE_MAX_DIM))
    out = io.BytesIO()
    img.save(out, format="JPEG", quality=85, optimize=True)
    await db.research_notes.update_one(
        {"id": note_id, "user_id": user.user_id},
        {"$set": {"image_data": out.getvalue(), "image_content_type": "image/jpeg", "has_image": True,
                   "updated_at": datetime.now(timezone.utc).isoformat()}},
    )
    return {"ok": True}


@api_router.get("/research/{note_id}/image")
async def get_research_note_image(note_id: str, user: User = Depends(get_current_user)):
    doc = await db.research_notes.find_one(
        {"id": note_id, "user_id": user.user_id}, {"_id": 0, "image_data": 1, "image_content_type": 1}
    )
    if not doc or not doc.get("image_data"):
        raise HTTPException(status_code=404, detail="Bilde ikke funnet")
    return Response(
        content=bytes(doc["image_data"]),
        media_type=doc.get("image_content_type") or "image/jpeg",
        headers={"Cache-Control": "private, max-age=86400"},
    )


# ─── Illustrators (public directory) ────────────────────────────────────────
class IllustratorCreate(BaseModel):
    name: str
    email: str
    portfolio_url: str
    style: str = ""
    services: str = ""
    # Honeypot field — bots often fill it; humans leave empty
    website: str = ""


@api_router.post("/illustrators")
async def create_illustrator(body: IllustratorCreate):
    # Honeypot spam guard — silently accept but don't store
    if (body.website or "").strip():
        return {"ok": True}

    name = (body.name or "").strip()
    email = (body.email or "").strip().lower()
    portfolio = (body.portfolio_url or "").strip()

    if not name or len(name) < 2:
        raise HTTPException(status_code=400, detail="Navn må fylles ut")
    if "@" not in email or "." not in email or len(email) < 5:
        raise HTTPException(status_code=400, detail="Ugyldig e-postadresse")
    if not portfolio or not (portfolio.startswith("http://") or portfolio.startswith("https://")):
        raise HTTPException(status_code=400, detail="Portfolio-lenke må begynne med http(s)://")

    doc = {
        "id": str(uuid.uuid4()),
        "edit_token": str(uuid.uuid4()),
        "name": name[:120],
        "email": email[:200],
        "portfolio_url": portfolio[:500],
        "style": (body.style or "").strip()[:600],
        "services": (body.services or "").strip()[:600],
        "is_public": True,
        "is_featured": False,
        "created_at": datetime.now(timezone.utc).isoformat(),
    }
    await db.illustrators.insert_one(doc)
    return {"ok": True, "id": doc["id"], "edit_token": doc["edit_token"]}


@api_router.get("/illustrators")
async def list_illustrators():
    """Public listing — no email is returned to the client. Featured listings sort first."""
    items = await db.illustrators.find(
        {"is_public": True},
        {"_id": 0, "email": 0},
    ).sort([("is_featured", -1), ("created_at", -1)]).to_list(200)
    for it in items:
        cover = await db.illustrator_images.find_one(
            {"illustrator_id": it["id"]}, {"_id": 0, "id": 1}, sort=[("order", 1)]
        )
        it["cover_image_id"] = cover["id"] if cover else None
    return items


ILLUSTRATOR_IMAGE_MAX_UPLOAD_BYTES = 8 * 1024 * 1024  # 8 MB raw upload cap
ILLUSTRATOR_IMAGE_MAX_DIM = 1400
ILLUSTRATOR_MAX_IMAGES = 6


def _encode_illustrator_image(data: bytes) -> bytes:
    if not data:
        raise HTTPException(status_code=400, detail="Tom fil")
    if len(data) > ILLUSTRATOR_IMAGE_MAX_UPLOAD_BYTES:
        raise HTTPException(status_code=400, detail="Bildet er for stort (maks 8 MB)")
    try:
        img = Image.open(io.BytesIO(data))
        img.load()
    except Exception:
        raise HTTPException(status_code=400, detail="Kunne ikke lese bildet — prøv en annen fil")

    if img.mode in ("RGBA", "LA", "P"):
        rgba = img.convert("RGBA")
        bg = Image.new("RGB", rgba.size, (255, 255, 255))
        bg.paste(rgba, mask=rgba.split()[-1])
        img = bg
    else:
        img = img.convert("RGB")

    img.thumbnail((ILLUSTRATOR_IMAGE_MAX_DIM, ILLUSTRATOR_IMAGE_MAX_DIM))
    out = io.BytesIO()
    img.save(out, format="JPEG", quality=85, optimize=True)
    return out.getvalue()


class IllustratorSelfUpdate(BaseModel):
    name: Optional[str] = None
    email: Optional[str] = None
    portfolio_url: Optional[str] = None
    style: Optional[str] = None
    services: Optional[str] = None


@api_router.get("/illustrators/edit/{edit_token}")
async def get_illustrator_for_edit(edit_token: str):
    """Illustrators have no accounts — this private edit_token (shown once at
    creation, never included in the public /illustrators listing) is what lets
    them come back later to update their listing or manage their image gallery."""
    doc = await db.illustrators.find_one(
        {"edit_token": edit_token}, {"_id": 0, "edit_token": 0}
    )
    if not doc:
        raise HTTPException(status_code=404, detail="Fant ikke oppføringen — sjekk lenken")
    doc["images"] = await db.illustrator_images.find(
        {"illustrator_id": doc["id"]}, {"_id": 0, "data": 0}
    ).sort("order", 1).to_list(ILLUSTRATOR_MAX_IMAGES)
    return doc


@api_router.patch("/illustrators/edit/{edit_token}")
async def update_illustrator_self(edit_token: str, body: IllustratorSelfUpdate):
    exists = await db.illustrators.find_one({"edit_token": edit_token}, {"_id": 0, "id": 1})
    if not exists:
        raise HTTPException(status_code=404, detail="Fant ikke oppføringen — sjekk lenken")

    updates: dict = {}
    if body.name is not None:
        name = body.name.strip()
        if len(name) < 2:
            raise HTTPException(status_code=400, detail="Navn må fylles ut")
        updates["name"] = name[:120]
    if body.email is not None:
        email = body.email.strip().lower()
        if "@" not in email or "." not in email or len(email) < 5:
            raise HTTPException(status_code=400, detail="Ugyldig e-postadresse")
        updates["email"] = email[:200]
    if body.portfolio_url is not None:
        portfolio = body.portfolio_url.strip()
        if not (portfolio.startswith("http://") or portfolio.startswith("https://")):
            raise HTTPException(status_code=400, detail="Portfolio-lenke må begynne med http(s)://")
        updates["portfolio_url"] = portfolio[:500]
    if body.style is not None:
        updates["style"] = body.style.strip()[:600]
    if body.services is not None:
        updates["services"] = body.services.strip()[:600]
    if not updates:
        raise HTTPException(status_code=400, detail="Ingenting å oppdatere")

    await db.illustrators.update_one({"edit_token": edit_token}, {"$set": updates})
    doc = await db.illustrators.find_one(
        {"edit_token": edit_token}, {"_id": 0, "edit_token": 0}
    )
    return doc


@api_router.post("/illustrators/edit/{edit_token}/images")
async def upload_illustrator_image(edit_token: str, file: UploadFile = File(...)):
    exists = await db.illustrators.find_one({"edit_token": edit_token}, {"_id": 0, "id": 1})
    if not exists:
        raise HTTPException(status_code=404, detail="Fant ikke oppføringen — sjekk lenken")
    illustrator_id = exists["id"]

    count = await db.illustrator_images.count_documents({"illustrator_id": illustrator_id})
    if count >= ILLUSTRATOR_MAX_IMAGES:
        raise HTTPException(status_code=400, detail=f"Maks {ILLUSTRATOR_MAX_IMAGES} bilder per oppføring")

    ct = (file.content_type or "").lower()
    if not ct.startswith("image/"):
        raise HTTPException(status_code=400, detail="Filen må være et bilde (PNG, JPG, WebP e.l.)")

    data = await file.read()
    jpeg_bytes = _encode_illustrator_image(data)

    doc = {
        "id": str(uuid.uuid4()),
        "illustrator_id": illustrator_id,
        "content_type": "image/jpeg",
        "data": jpeg_bytes,
        "order": count,
        "created_at": datetime.now(timezone.utc).isoformat(),
    }
    await db.illustrator_images.insert_one(doc)
    return {"ok": True, "id": doc["id"]}


@api_router.delete("/illustrators/edit/{edit_token}/images/{image_id}")
async def delete_illustrator_image(edit_token: str, image_id: str):
    exists = await db.illustrators.find_one({"edit_token": edit_token}, {"_id": 0, "id": 1})
    if not exists:
        raise HTTPException(status_code=404, detail="Fant ikke oppføringen — sjekk lenken")
    r = await db.illustrator_images.delete_one({"id": image_id, "illustrator_id": exists["id"]})
    if r.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Bilde ikke funnet")
    return {"ok": True}


@api_router.get("/illustrators/{illustrator_id}/images")
async def list_illustrator_images(illustrator_id: str):
    return await db.illustrator_images.find(
        {"illustrator_id": illustrator_id}, {"_id": 0, "data": 0}
    ).sort("order", 1).to_list(ILLUSTRATOR_MAX_IMAGES)


@api_router.get("/illustrators/{illustrator_id}/images/{image_id}")
async def get_illustrator_image(illustrator_id: str, image_id: str):
    doc = await db.illustrator_images.find_one(
        {"id": image_id, "illustrator_id": illustrator_id}, {"_id": 0, "data": 1, "content_type": 1}
    )
    if not doc:
        raise HTTPException(status_code=404, detail="Bilde ikke funnet")
    return Response(
        content=bytes(doc["data"]),
        media_type=doc.get("content_type") or "image/jpeg",
        headers={"Cache-Control": "public, max-age=86400"},
    )


class IllustratorCheckoutRequest(BaseModel):
    illustrator_id: str
    origin_url: str


@api_router.post("/illustrators/checkout")
async def illustrator_checkout(body: IllustratorCheckoutRequest):
    """Guest checkout for the optional 'fremhevet' (featured) upgrade — illustrators
    don't have accounts, so we identify the listing by id and email Stripe directly."""
    illustrator = await db.illustrators.find_one({"id": body.illustrator_id}, {"_id": 0})
    if not illustrator:
        raise HTTPException(status_code=404, detail="Oppføring ikke funnet")

    prices = stripe.Price.list(lookup_keys=["bragr_illustrator_featured_nok"], active=True, limit=1).data
    if not prices:
        raise HTTPException(status_code=500, detail="Pris ikke funnet")
    price = prices[0]

    origin_ascii = _ascii_url(body.origin_url)

    session = stripe.checkout.Session.create(
        customer_email=illustrator["email"],
        line_items=[{"price": price.id, "quantity": 1}],
        mode="subscription",
        success_url=f"{origin_ascii}/illustratorer?fremhevet=1",
        cancel_url=f"{origin_ascii}/illustratorer",
        metadata={"illustrator_id": body.illustrator_id},
        managed_payments={"enabled": True},
        subscription_data={"metadata": {"illustrator_id": body.illustrator_id}},
        automatic_tax={"enabled": False},
        adaptive_pricing={"enabled": False},
    )
    return {"checkout_url": session.url}



@api_router.post("/characters/extract")
async def extract_characters(user: User = Depends(get_current_user)):
    """Read all scenes and ask Claude to extract character profiles."""
    if not ANTHROPIC_API_KEY:
        raise HTTPException(status_code=503, detail="AI ikke tilgjengelig")
    scenes = await db.scenes.find({"user_id": user.user_id}, {"_id": 0}).sort("order", 1).to_list(500)
    if not scenes:
        raise HTTPException(status_code=400, detail="Legg inn scener i /manuskript først")
    # Build a compact corpus
    parts = []
    for s in scenes[:30]:
        c = (s.get("content") or "").strip()
        if c:
            parts.append(f"[{s.get('title', 'Scene')}] {c[:2000]}")
    corpus = "\n\n---\n\n".join(parts)[:15000]
    system = (
        "Du er en litterær redaktør som identifiserer karakterer i romaner. "
        "Trekk ut alle navngitte personer og bygg psykologiske profiler basert på hva "
        "teksten faktisk viser — ikke gjett. Svar KUN i gyldig JSON."
    )
    prompt = (
        "Analyser teksten under. Returner JSON i formen:\n"
        '{"characters": [{"name":"", "role":"", "appearance":"", "inner_struggle":"", '
        '"outer_struggle":"", "relationships":"", "arc":"", "voice_notes":""}]}\n\n'
        "Skriv beskrivelsene på norsk. Vær konkret og tekstnær. Ikke oppfinn.\n\n"
        f"TEKST:\n{corpus}"
    )
    try:
        parts = []
        async for chunk in stream_llm_text(*BRAGR_MODEL_ANALYSIS, system, prompt, ANTHROPIC_API_KEY, temperature=0.4):
            parts.append(chunk)
        raw = "".join(parts).strip()
        m = re.search(r"\{.*\}", raw, re.DOTALL)
        data = json.loads(m.group(0) if m else raw)
        chars = data.get("characters", []) or []
    except Exception as e:
        logger.warning("character extract failed: %s", e)
        raise HTTPException(status_code=500, detail="AI-uttrekk feilet, prøv igjen")
    # Upsert each: match by (user_id, name) case-insensitive
    saved = []
    for c in chars:
        name = (c.get("name") or "").strip()
        if not name:
            continue
        existing = await db.characters.find_one(
            {"user_id": user.user_id, "name": {"$regex": f"^{re.escape(name)}$", "$options": "i"}},
            {"_id": 0},
        )
        payload = {
            "user_id": user.user_id,
            "name": name,
            "role": (c.get("role") or "").strip(),
            "appearance": (c.get("appearance") or "").strip(),
            "inner_struggle": (c.get("inner_struggle") or "").strip(),
            "outer_struggle": (c.get("outer_struggle") or "").strip(),
            "relationships": (c.get("relationships") or "").strip(),
            "arc": (c.get("arc") or "").strip(),
            "voice_notes": (c.get("voice_notes") or "").strip(),
            "updated_at": datetime.now(timezone.utc).isoformat(),
        }
        if existing:
            await db.characters.update_one({"id": existing["id"], "user_id": user.user_id}, {"$set": payload})
            saved.append({**existing, **payload})
        else:
            payload.update({"id": str(uuid.uuid4()), "created_at": payload["updated_at"]})
            await db.characters.insert_one(payload)
            payload.pop("_id", None)
            saved.append(payload)
    return {"count": len(saved), "characters": saved}


class WritingGoal(BaseModel):
    total_goal: int = 0
    session_goal: int = 0


@api_router.get("/manuscript/goals")
async def get_goals(user: User = Depends(get_current_user)):
    doc = await db.writing_goals.find_one({"user_id": user.user_id}, {"_id": 0}) or {"total_goal": 0, "session_goal": 0}
    return doc


@api_router.put("/manuscript/goals")
async def set_goals(body: WritingGoal, user: User = Depends(get_current_user)):
    await db.writing_goals.update_one(
        {"user_id": user.user_id},
        {"$set": {
            "user_id": user.user_id,
            "total_goal": max(0, body.total_goal),
            "session_goal": max(0, body.session_goal),
            "updated_at": datetime.now(timezone.utc).isoformat(),
        }},
        upsert=True,
    )
    return {"total_goal": body.total_goal, "session_goal": body.session_goal}


@api_router.get("/manuscript/compile.docx")
async def compile_manuscript_docx(user: User = Depends(get_current_user)):
    scenes = await db.scenes.find({"user_id": user.user_id}, {"_id": 0}).sort("order", 1).to_list(1000)
    doc = Document()
    # Base style — Times New Roman 12pt, double spacing (industry standard for novel manuscripts)
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Title page
    t = doc.add_paragraph()
    t.alignment = 1  # center
    r = t.add_run(user.name or "Manuskript")
    r.font.size = Pt(24)
    r.bold = True
    doc.add_paragraph()
    doc.add_paragraph()

    if not scenes:
        doc.add_paragraph("(Ingen scener enda.)")
    for i, s in enumerate(scenes):
        # Page break between scenes (except first)
        if i > 0:
            doc.add_page_break()
        heading = doc.add_paragraph()
        heading.alignment = 1
        hr = heading.add_run(s.get("title") or f"Scene {i+1}")
        hr.font.size = Pt(16)
        hr.bold = True
        if s.get("synopsis"):
            syn = doc.add_paragraph()
            sr = syn.add_run(s["synopsis"])
            sr.italic = True
            sr.font.size = Pt(11)
        doc.add_paragraph()
        content = (s.get("content") or "").strip()
        for para in content.split("\n"):
            p = doc.add_paragraph(para)
            p.paragraph_format.first_line_indent = Cm(1)
            p.paragraph_format.line_spacing = 2.0

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    filename = f"bragarmaal-manuskript-{datetime.now(timezone.utc).strftime('%Y-%m-%d')}.docx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ---------- Snapshots (scene version history) ----------
class SnapshotCreate(BaseModel):
    scene_id: str
    label: Optional[str] = ""


@api_router.post("/manuscript/{scene_id}/snapshots")
async def create_snapshot(scene_id: str, body: SnapshotCreate, user: User = Depends(get_current_user)):
    scene = await db.scenes.find_one({"id": scene_id, "user_id": user.user_id}, {"_id": 0})
    if not scene:
        raise HTTPException(status_code=404, detail="Scene ikke funnet")
    snap = {
        "id": str(uuid.uuid4()),
        "user_id": user.user_id,
        "scene_id": scene_id,
        "title": scene.get("title", ""),
        "content": scene.get("content", ""),
        "word_count": scene.get("word_count", 0),
        "label": (body.label or "").strip(),
        "created_at": datetime.now(timezone.utc).isoformat(),
    }
    await db.snapshots.insert_one(snap)
    snap.pop("_id", None)
    return snap


@api_router.get("/manuscript/{scene_id}/snapshots")
async def list_snapshots(scene_id: str, user: User = Depends(get_current_user)):
    docs = await db.snapshots.find(
        {"scene_id": scene_id, "user_id": user.user_id}, {"_id": 0}
    ).sort("created_at", -1).to_list(100)
    return docs


@api_router.post("/manuscript/snapshots/{snapshot_id}/restore")
async def restore_snapshot(snapshot_id: str, user: User = Depends(get_current_user)):
    snap = await db.snapshots.find_one({"id": snapshot_id, "user_id": user.user_id}, {"_id": 0})
    if not snap:
        raise HTTPException(status_code=404, detail="Øyeblikksbilde ikke funnet")
    await db.scenes.update_one(
        {"id": snap["scene_id"], "user_id": user.user_id},
        {"$set": {
            "title": snap.get("title", ""),
            "content": snap.get("content", ""),
            "word_count": snap.get("word_count", 0),
            "updated_at": datetime.now(timezone.utc).isoformat(),
        }},
    )
    doc = await db.scenes.find_one({"id": snap["scene_id"], "user_id": user.user_id}, {"_id": 0})
    return doc


@api_router.delete("/manuscript/snapshots/{snapshot_id}")
async def delete_snapshot(snapshot_id: str, user: User = Depends(get_current_user)):
    r = await db.snapshots.delete_one({"id": snapshot_id, "user_id": user.user_id})
    if r.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Ikke funnet")
    return {"ok": True}


# ---------- Custom AI helpers (user's own key + persona) ----------
HELPER_PROVIDERS = {"openai", "anthropic", "gemini"}


class HelperCreate(BaseModel):
    name: str
    provider: str
    model_id: str
    api_key: str
    persona_addon: Optional[str] = ""


class HelperUpdate(BaseModel):
    name: Optional[str] = None
    provider: Optional[str] = None
    model_id: Optional[str] = None
    api_key: Optional[str] = None
    persona_addon: Optional[str] = None


def _sanitize_helper(doc: dict) -> dict:
    d = dict(doc)
    key = d.get("api_key", "")
    d["api_key_preview"] = ("…" + key[-4:]) if key and len(key) >= 4 else ""
    d.pop("api_key", None)
    d.pop("_id", None)
    return d


@api_router.get("/helpers")
async def list_helpers(user: User = Depends(get_current_user)):
    docs = await db.helpers.find(
        {"user_id": user.user_id}, {"_id": 0}
    ).sort("created_at", 1).to_list(50)
    return [_sanitize_helper(d) for d in docs]


@api_router.post("/helpers")
async def create_helper(body: HelperCreate, user: User = Depends(get_current_user)):
    if body.provider not in HELPER_PROVIDERS:
        raise HTTPException(status_code=400, detail="Ukjent leverandør. Bruk openai, anthropic eller gemini.")
    if not body.name.strip() or not body.model_id.strip() or not body.api_key.strip():
        raise HTTPException(status_code=400, detail="Navn, modell og API-nøkkel er påkrevd")
    doc = {
        "id": str(uuid.uuid4()),
        "user_id": user.user_id,
        "name": body.name.strip()[:80],
        "provider": body.provider,
        "model_id": body.model_id.strip()[:120],
        "api_key": body.api_key.strip(),
        "persona_addon": (body.persona_addon or "").strip()[:3000],
        "created_at": datetime.now(timezone.utc).isoformat(),
    }
    await db.helpers.insert_one(doc)
    return _sanitize_helper(doc)


@api_router.patch("/helpers/{helper_id}")
async def update_helper(helper_id: str, body: HelperUpdate, user: User = Depends(get_current_user)):
    existing = await db.helpers.find_one({"id": helper_id, "user_id": user.user_id}, {"_id": 0})
    if not existing:
        raise HTTPException(status_code=404, detail="Ikke funnet")
    updates = {}
    if body.name is not None: updates["name"] = body.name.strip()[:80]
    if body.model_id is not None: updates["model_id"] = body.model_id.strip()[:120]
    if body.provider is not None:
        if body.provider not in HELPER_PROVIDERS:
            raise HTTPException(status_code=400, detail="Ukjent leverandør")
        updates["provider"] = body.provider
    if body.api_key is not None and body.api_key.strip():
        updates["api_key"] = body.api_key.strip()
    if body.persona_addon is not None:
        updates["persona_addon"] = body.persona_addon.strip()[:3000]
    if updates:
        await db.helpers.update_one({"id": helper_id, "user_id": user.user_id}, {"$set": updates})
    doc = await db.helpers.find_one({"id": helper_id, "user_id": user.user_id}, {"_id": 0})
    return _sanitize_helper(doc)


@api_router.delete("/helpers/{helper_id}")
async def delete_helper(helper_id: str, user: User = Depends(get_current_user)):
    r = await db.helpers.delete_one({"id": helper_id, "user_id": user.user_id})
    if r.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Ikke funnet")
    return {"ok": True}


# ---------- Voice profile ----------
@api_router.post("/voice/analyze")
async def analyze_user_voice(user: User = Depends(get_current_user)):
    all_samples = await db.samples.find({"user_id": user.user_id}, {"_id": 0}).to_list(500)
    if not all_samples:
        raise HTTPException(status_code=400, detail="Legg til minst én skriveprøve først")

    # Base the voice profile on HUMAN samples only (exclude hybrid/pure AI)
    human_samples = [s for s in all_samples if s.get("category", "ren_menneske_ny") in HUMAN_CATEGORIES]
    if not human_samples:
        raise HTTPException(
            status_code=400,
            detail="Ingen 'ren menneske'-prøver funnet. Hybrid- og AI-prøver brukes ikke som basis for stemmen."
        )

    stats = analyze_voice(human_samples)
    ai_stuff = await build_style_summary(human_samples, stats)

    profile = {
        "user_id": user.user_id,
        "updated_at": datetime.now(timezone.utc).isoformat(),
        "total_samples": len(human_samples),
        **stats,
        **ai_stuff,
    }
    await db.voice_profiles.update_one(
        {"user_id": user.user_id},
        {"$set": profile},
        upsert=True,
    )
    return profile


@api_router.get("/voice/profile")
async def get_voice_profile(user: User = Depends(get_current_user)):
    doc = await db.voice_profiles.find_one({"user_id": user.user_id}, {"_id": 0})
    if not doc:
        return None
    return doc


VOICE_SCOPE_MIN_WORDS = 40


class VoiceAnalyzeScopeBody(BaseModel):
    scope: Literal["scene", "sample", "text"]
    id: Optional[str] = None  # scene_id or sample_id — required for scope in {scene, sample}
    text: Optional[str] = None  # required for scope == text


@api_router.post("/voice/analyze-scope")
async def analyze_voice_scope(body: VoiceAnalyzeScopeBody, user: User = Depends(get_current_user)):
    """
    Kjør samme stemmeanalyse som /voice/analyze, men kun på én enkelt enhet
    (én scene, én prøve, eller et limt inn avsnitt) i stedet for alle prøver
    samlet. Resultatet lagres IKKE som brukerens hovedprofil i voice_profiles —
    det er en engangsrapport for den valgte enheten.
    """
    label = ""
    if body.scope == "scene":
        if not body.id:
            raise HTTPException(status_code=400, detail="Mangler scene-id")
        scene = await db.scenes.find_one({"id": body.id, "user_id": user.user_id}, {"_id": 0})
        if not scene:
            raise HTTPException(status_code=404, detail="Scene ikke funnet")
        content = scene.get("content", "")
        label = scene.get("title") or "Uten tittel"
    elif body.scope == "sample":
        if not body.id:
            raise HTTPException(status_code=400, detail="Mangler prøve-id")
        sample = await db.samples.find_one({"id": body.id, "user_id": user.user_id}, {"_id": 0})
        if not sample:
            raise HTTPException(status_code=404, detail="Prøve ikke funnet")
        content = sample.get("content", "")
        label = sample.get("title") or "Uten tittel"
    else:  # text
        content = (body.text or "").strip()
        label = "Limt inn avsnitt"

    word_count = _word_count(content)
    if word_count < VOICE_SCOPE_MIN_WORDS:
        raise HTTPException(
            status_code=400,
            detail=f"For lite tekst for en pålitelig analyse — minst {VOICE_SCOPE_MIN_WORDS} ord, denne enheten har {word_count}."
        )

    unit = [{"content": content}]
    stats = analyze_voice(unit)
    ai_stuff = await build_style_summary(unit, stats)

    return {
        "scope": body.scope,
        "scope_label": label,
        "word_count": word_count,
        **stats,
        **ai_stuff,
    }


# ---------- Generation ----------
def _anti_slop_rules() -> str:
    """
    Kjerneregler mot AI-slop og «corporate-språk».
    Brukes av /generate, humanisering og «Skriv om i min stemme»-verktøyene.
    Norsk-spesifikk: fanger opp de faktiske ordene og strukturene som avslører AI-bruk på norsk.
    """
    return (
        "\n\n=== ANTI-SLOP-REGLER (må følges strengt) ===\n"
        "\n"
        "── A. FORBUDTE NORSKE FRASER (dette er norsk AI-signatur) ──\n"
        "«dykk ned i», «la oss se nærmere på», «i en verden der», «i dagens (raske / digitale / travle) verden»,\n"
        "«i kjernen av», «i essensen av», «til syvende og sist», «når alt kommer til alt», «i lys av»,\n"
        "«det er verdt å merke seg», «det er viktig å forstå», «det bør bemerkes», «det kan sies at»,\n"
        "«spiller en sentral rolle», «utgjør en viktig faktor», «har stor betydning for»,\n"
        "«danner grunnlaget for», «peker på viktigheten av», «legger til rette for», «tar i bruk»,\n"
        "«utnytte fordelene», «effektivisere prosessen», «på en effektiv måte», «i en travel hverdag»,\n"
        "«en rekke fordeler», «en rekke muligheter», «en rekke utfordringer»,\n"
        "«ikke bare … men også», «både … og …» (overbrukt),\n"
        "«sammenvevd», «sømløst», «robust», «kraftfull», «grunnleggende», «betydningsfull»,\n"
        "«orkestrere», «revolusjonere», «låse opp potensialet», «utnytte muligheter»,\n"
        "«dette handler ikke bare om X — det handler om Y», «det er ikke X. Det er Y.» (AI-tvillingpar).\n"
        "\n"
        "── B. FORBUDTE OVERGANGER ──\n"
        "Ikke start setninger med: «Videre», «Dessuten», «Sammenfattet», «For det første/andre/tredje»,\n"
        "«Til slutt», «For å oppsummere», «Alt i alt», «I det store og hele», «På en annen side».\n"
        "Engelske ekvivalenter er også forbudt: furthermore, moreover, additionally, in conclusion,\n"
        "delve, realm, tapestry, testament, seamlessly, leverage, harness, orchestrate.\n"
        "\n"
        "── C. BYTT OPPBLÅSTE VERB MED HVERDAGSLIGE ──\n"
        "«utnytte» → «bruke». «anvende» → «bruke». «gjennomføre» → «gjøre».\n"
        "«fasilitere» → «hjelpe». «kommunisere» → «si» eller «snakke». «demonstrere» → «vise».\n"
        "«illustrere» → «vise». «konstatere» → «se». «erverve» → «få». «avvikle» → «avslutte».\n"
        "«besitte» → «ha». «vedrørende» → «om». «i forbindelse med» → «med». «angående» → «om».\n"
        "«implementere» → «lage / sette i gang». «optimalisere» → «gjøre bedre».\n"
        "\n"
        "── D. RYTME OG SETNINGSSTRUKTUR (viktigst!) ──\n"
        "1. Bland setningslengder ujevnt. Veldig korte setninger. Så en lang, flerdelt en som beveger seg,\n"
        "   snubler, tar en pause — og lander. Så en kort igjen.\n"
        "2. Unngå «tre-regelen»: AI lister opp i grupper på tre (rytme, tone, ordvalg). Bryt mønsteret —\n"
        "   bruk to, eller fire, eller en enkelt tydelig.\n"
        "3. Ikke over-forklar. Kutt oppsummeringer på slutten av avsnitt. Kutt meta-kommentarer\n"
        "   («som nevnt», «som vi har sett», «dette viser at»).\n"
        "4. Unngå passiv der aktiv er mulig. «Det ble bestemt» → «Hun bestemte».\n"
        "5. Tillat enkeltordssetninger. Slik. Fordi de bærer.\n"
        "6. Bruk «og» og «men» som setningsstartere når rytmen ber om det. AI unngår dette — mennesker gjør det.\n"
        "7. Ellipser og ufullstendige setninger er ok — som når man tenker seg om.\n"
        "\n"
        "── E. HVERDAGSNORSK (byggeklosser for menneskelig fortellerstemme) ──\n"
        "– «Man» sparsommelig. Direkte «jeg», «du», «vi» der det passer.\n"
        "– Konkrete sanselige detaljer i stedet for abstrakter. «Kaffen ble kald» ikke «en observasjon av kvalitetsmessig forringelse».\n"
        "– Talespråklige innskudd der forfatterens stemme viser det: «altså», «liksom», «jo», «vel», «kanskje».\n"
        "– Sammentrekninger og korthet: «jeg er ikke» kan bli «jeg er´kke» hvis stemmen tåler det.\n"
        "– Norske idiomer og folkelige vendinger som forfatteren faktisk bruker (ikke oppfinn nye).\n"
        "– Unngå lange sammensatte substantiv («kvalitetsforbedringstiltak») — de skriker byråkrati.\n"
        "\n"
        "── F. MENNESKELIG LOGIKK ──\n"
        "1. Legg til perspektiv. Mennesker kobler og tolker — vi lister ikke bare fakta.\n"
        "2. Aksepter friksjon. La ideer bevege seg naturlig, uten å begrunne hvert steg.\n"
        "3. Litt ubalanse er menneskelig. En liten selvmotsigelse, en tanke som skifter retning,\n"
        "   et poeng som ikke blir helt fulgt opp — det er slik ekte skrift beveger seg.\n"
        "4. Vær presis der det gjelder, sløv der det ikke gjør det. AI er alltid presis — mennesker vekter.\n"
        "\n"
        "── G. LEVERANSE ──\n"
        "Ikke skriv «her er teksten:». Ingen preamble. Ingen «håper dette hjelper». Ingen etterrasjonalisering.\n"
        "Bare selve teksten. Punktum.\n"
    )


def build_voice_system_prompt(profile: Optional[dict], samples: List[dict], humanize_level: int) -> str:
    base = (
        "Du er en skygge-skriver som gjenskaper en spesifikk norsk forfatters stemme perfekt. "
        "Du skriver ALLTID på norsk (bokmål med mindre prøvene tydelig er nynorsk). "
        "Målet er tekst som verken føles maskinell eller generisk. "
        "Sluttresultatet skal alltid være brukerens egen stemme, slik den fremgår av hennes "
        "prøvetekster — aldri en imitasjon av noen andre."
    )

    base += _anti_slop_rules()

    if profile:
        base += "\n--- FORFATTERENS STEMME ---\n"
        if profile.get("tone_description"):
            base += f"Tone: {profile['tone_description']}\n"
        if profile.get("style_summary"):
            base += f"Stil: {profile['style_summary']}\n"
        if profile.get("signature_phrases"):
            base += f"Signaturord/uttrykk å nikke til (bruk sparsomt, ikke tvinge inn): {', '.join(profile['signature_phrases'])}\n"
        if profile.get("avg_sentence_length"):
            base += f"Gjennomsnittlig setningslengde: {profile['avg_sentence_length']} ord. Etterlign dette omtrentlig, men varier bevisst.\n"

    # Include short excerpts from samples for style anchoring
    if samples:
        excerpts = []
        total = 0
        for s in samples[:5]:
            snippet = s["content"][:900]
            excerpts.append(snippet)
            total += len(snippet)
            if total > 4000:
                break
        base += "\n--- REFERANSE FRA FORFATTERENS EGNE TEKSTER (etterlign rytme og ordvalg) ---\n"
        base += "\n\n---\n\n".join(excerpts)

    if humanize_level >= 2:
        base += (
            "\n\nEkstra humanisering: Legg inn små naturlige uregelmessigheter — "
            "en ufullstendig setning her og der, en spontan tanke i parentes, "
            "en dagligdags vending. Ikke overdriv. Tekst skal ikke se 'redigert' ut."
        )
    if humanize_level >= 3:
        base += (
            "\n\nMaksimal humanisering: Skriv som førsteutkast fra en menneskelig forfatter. "
            "Bruk 'og' og 'men' som setningsstartere når det passer. Tillat gjentakelser "
            "hvis de forsterker rytmen. Ingen glatt AI-flyt."
        )
    return base


class GenerateBody(BaseModel):
    mode: Literal["prompt", "continue", "humanize", "next_steps", "reflect"]
    text: str
    # Bragarmål velger automatisk beste modell for skriving (Claude Sonnet 4.6).
    # Denne kan fortsatt overstyres av interne kall (f.eks. helper:<id> BYOK).
    model: str = "claude-sonnet-4-6"
    humanize_level: int = 1  # 1..3
    length: Literal["kort", "medium", "lang"] = "medium"
    temperature: float = 0.7  # 0.2..1.2 — kun brukt i mode "next_steps"


@api_router.post("/generate")
async def generate(body: GenerateBody, user: User = Depends(get_current_user)):
    # Auto-grant beta if within first N users, then check subscription
    await ensure_beta_flag(user.user_id)
    sub = await get_user_subscription_status(user.user_id)
    if not sub["active"]:
        raise HTTPException(
            status_code=402,
            detail="Ditt medlemskap er ikke aktivt. Åpne betalingssiden for å fortsette å generere tekst.",
        )

    # Resolve model: either built-in or user's own AI helper (helper:<id>)
    helper_persona = ""
    helper_name = ""
    if body.model.startswith("helper:"):
        helper_id = body.model.split(":", 1)[1]
        helper = await db.helpers.find_one({"id": helper_id, "user_id": user.user_id}, {"_id": 0})
        if not helper:
            raise HTTPException(status_code=404, detail="AI-hjelper ikke funnet")
        provider = helper["provider"]
        model = helper["model_id"]
        api_key_to_use = helper["api_key"]
        helper_persona = helper.get("persona_addon", "") or ""
        helper_name = helper.get("name", "")
    else:
        if body.model not in ALLOWED_MODELS:
            raise HTTPException(status_code=400, detail="Ukjent modell")
        provider, model = ALLOWED_MODELS[body.model]
        api_key_to_use = PROVIDER_API_KEYS.get(provider)
        if not api_key_to_use:
            raise HTTPException(status_code=503, detail="KI-generering er ikke tilgjengelig på denne serveren ennå")

    profile = await db.voice_profiles.find_one({"user_id": user.user_id}, {"_id": 0})
    all_samples = await db.samples.find({"user_id": user.user_id}, {"_id": 0}).sort("created_at", -1).to_list(30)
    # Only use HUMAN samples as style anchors for generation
    samples = [s for s in all_samples if s.get("category", "ren_menneske_ny") in HUMAN_CATEGORIES][:20]

    system = build_voice_system_prompt(profile, samples, body.humanize_level)

    if helper_persona:
        system += (
            f"\n\n--- BRUKERENS EGEN AI-HJELPER: «{helper_name}» ---\n"
            "Brukeren har trent opp denne AI-en med en spesifikk arbeids- og tonemåte over tid. "
            "Nedenstående instruks er lagt til av brukeren. Følg den så lenge den ikke bryter med "
            "kravene om å bevare brukerens egen stemme fra prøvetekstene:\n"
            f"{helper_persona}\n"
        )

    length_hint = {
        "kort": "Skriv omtrent 80-150 ord.",
        "medium": "Skriv omtrent 200-350 ord.",
        "lang": "Skriv omtrent 450-700 ord.",
    }[body.length]

    if body.mode == "prompt":
        user_msg = f"{length_hint}\n\nSkriv en tekst basert på dette utgangspunktet, helt i forfatterens stemme:\n\n{body.text}"
    elif body.mode == "continue":
        user_msg = f"{length_hint}\n\nFortsett følgende tekst sømløst i samme stemme og rytme. Ikke gjenta det som allerede står. Start rett der teksten slipper:\n\n---\n{body.text}\n---"
    elif body.mode == "next_steps":
        user_msg = (
            "Du er en litterær sparringspartner — ikke en skriver.\n\n"
            "Under er en tekst forfatteren jobber med. Skriv IKKE ferdig prosa. Skriv IKKE et neste avsnitt.\n\n"
            "Din oppgave: Gi 3–5 korte, konkrete forslag til hvor teksten kan gå videre. Retninger, uventede vinkler, "
            "sanselige detaljer å prøve, spørsmål å utforske, motiv-linjer å strekke. "
            "Presenter som en nummerert liste (1., 2., 3.). Hvert punkt: én eller to setninger. "
            "Snakk direkte til forfatteren («du»). Ikke forklar, ikke oppsummer, ikke skriv preamble. "
            "Vær konkret — ikke generiske råd som «utdyp følelsen». Grip fatt i noe faktisk i teksten.\n\n"
            f"---\n{body.text}\n---"
        )
    elif body.mode == "reflect":
        user_msg = (
            "Du er en litterær sparringspartner — ikke en redigerer.\n\n"
            "Les teksten under. Skriv IKKE forslag til endringer, IKKE ferdig prosa, IKKE en liste over ting "
            "som kan forbedres. Skriv en kort editorisk lesning: hva du la merke til (bilder, rytme, tone), "
            "hvor spenningen sitter, hva som er ubesvart, hva stemmen prøver å fortelle deg. "
            "Personlig og nøktern — ikke overdriv, ikke smiger. 4–7 setninger, én paragraf.\n\n"
            "Ikke gi råd — bare speil hva du så. Slik en god venn som leser bra kunne gjort det.\n\n"
            f"---\n{body.text}\n---"
        )
    else:  # humanize
        user_msg = (
            "Skriv følgende tekst på nytt slik at den fremstår helt menneskelig og i forfatterens stemme. "
            "Fjern AI-signaturer, bryt opp glatt flyt, behold meningen. Lever kun den omskrevne teksten:\n\n"
            f"---\n{body.text}\n---"
        )

    gen_temperature = max(0.2, min(1.2, float(body.temperature or 0.7))) if body.mode == "next_steps" else 0.7

    async def stream_llm():
        try:
            async for chunk in stream_llm_text(provider, model, system, user_msg, api_key_to_use, temperature=gen_temperature):
                yield f"data: {json.dumps({'delta': chunk})}\n\n"
            yield f"data: {json.dumps({'done': True})}\n\n"
        except Exception as e:
            logger.exception("Generation error")
            yield f"data: {json.dumps({'error': str(e)})}\n\n"

    async def stream_xai():
        if not XAI_API_KEY:
            yield f"data: {json.dumps({'error': 'xAI-nøkkel mangler'})}\n\n"
            return
        try:
            client = AsyncOpenAI(api_key=XAI_API_KEY, base_url="https://api.x.ai/v1")
            resp = await client.chat.completions.create(
                model=model,
                stream=True,
                messages=[
                    {"role": "system", "content": system},
                    {"role": "user", "content": user_msg},
                ],
            )
            async for chunk in resp:
                delta = chunk.choices[0].delta.content if chunk.choices else None
                if delta:
                    yield f"data: {json.dumps({'delta': delta})}\n\n"
            yield f"data: {json.dumps({'done': True})}\n\n"
        except Exception as e:
            logger.exception("xAI generation error")
            yield f"data: {json.dumps({'error': str(e)})}\n\n"

    stream_fn = stream_xai if provider == "xai" else stream_llm
    return StreamingResponse(
        stream_fn(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


# ---------- Sporede endringer (diff-basert redigering) ----------
class TrackEditBody(BaseModel):
    text: str
    focus: Literal["korrektur", "flyt", "begge"] = "begge"


def _diff_segments(original: str, edited: str) -> list:
    """Word-level diff between original and AI-edited text, as segments the
    frontend can render inline with per-change accept/reject."""
    a_tokens = re.findall(r"\s+|\S+", original)
    b_tokens = re.findall(r"\s+|\S+", edited)
    sm = difflib.SequenceMatcher(None, a_tokens, b_tokens, autojunk=False)
    segments = []
    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        if tag == "equal":
            segments.append({"type": "equal", "text": "".join(a_tokens[i1:i2])})
        else:
            segments.append({
                "type": "change",
                "original": "".join(a_tokens[i1:i2]),
                "replacement": "".join(b_tokens[j1:j2]),
            })
    return segments


@api_router.post("/edit/track")
async def track_edit(body: TrackEditBody, user: User = Depends(get_current_user)):
    await ensure_beta_flag(user.user_id)
    sub = await get_user_subscription_status(user.user_id)
    if not sub["active"]:
        raise HTTPException(
            status_code=402,
            detail="Ditt medlemskap er ikke aktivt. Åpne betalingssiden for å fortsette.",
        )

    text = body.text.strip()
    if not text:
        raise HTTPException(status_code=400, detail="Tom tekst")
    if len(text) > 12000:
        raise HTTPException(
            status_code=400,
            detail="Teksten er for lang for sporede endringer (maks ca. 2000 ord). Prøv et kortere utdrag.",
        )
    if not ANTHROPIC_API_KEY:
        raise HTTPException(status_code=503, detail="KI-redigering er ikke tilgjengelig på denne serveren ennå")

    profile = await db.voice_profiles.find_one({"user_id": user.user_id}, {"_id": 0})
    all_samples = await db.samples.find({"user_id": user.user_id}, {"_id": 0}).sort("created_at", -1).to_list(30)
    samples = [s for s in all_samples if s.get("category", "ren_menneske_ny") in HUMAN_CATEGORIES][:20]
    system = build_voice_system_prompt(profile, samples, humanize_level=1)
    system += (
        "\n\n--- OPPGAVE: SPOREDE ENDRINGER ---\n"
        "Du er nå en varsom korrekturleser, ikke en gjenskriver. Du får en tekst forfatteren "
        "allerede har skrevet i sin egen stemme. Foreslå kun presise, lokale endringer — "
        "grammatikk, tegnsetting, ordgjentakelser, klossete setninger, åpenbare skrivefeil, "
        "og steder der flyten butter. IKKE omskriv setninger som allerede fungerer. IKKE "
        "endre stemmen, ordvalget eller tonen. IKKE legg til eller fjern innhold eller mening. "
        "Behold linjeskift og avsnittsinndeling identisk. Endre minst mulig — helst under "
        "10-20% av teksten. Svar KUN med den redigerte teksten — ingen kommentarer, ingen "
        "forklaring, ingen anførselstegn rundt."
    )
    focus_hint = {
        "korrektur": "Fokuser kun på ren korrektur: skrivefeil, grammatikk, tegnsetting.",
        "flyt": "Fokuser på flyt og rytme: klossete setninger, gjentakelser, tempo — ikke ren rettskriving.",
        "begge": "Se etter både korrektur og flyt.",
    }[body.focus]
    user_msg = f"{focus_hint}\n\n---\n{text}\n---"

    try:
        parts = []
        async for chunk in stream_llm_text(
            *BRAGR_MODEL_ANALYSIS, system, user_msg, ANTHROPIC_API_KEY, temperature=0.3, max_tokens=6000,
        ):
            parts.append(chunk)
        edited = "".join(parts).strip()
    except Exception:
        logger.exception("Track-edit generation error")
        raise HTTPException(status_code=502, detail="Kunne ikke generere endringsforslag — prøv igjen")

    if not edited:
        raise HTTPException(status_code=502, detail="Fikk ikke noe svar — prøv igjen")

    return {"segments": _diff_segments(text, edited)}


# ---------- AI detection heuristic ----------
class DetectBody(BaseModel):
    text: str


AI_MARKERS = [
    r"\bi en verden der\b",
    r"\bla oss dykke\b",
    r"\btil syvende og sist\b",
    r"\bi lys av\b",
    r"\bdet er viktig å merke seg\b",
    r"\bnår alt kommer til alt\b",
    r"\bdet er verdt å nevne\b",
    r"\bfor å oppsummere\b",
    r"\bfør vi går videre\b",
    r"\bsammenfattet\b",
    r"\bfremtidens\b",
    r"\bi en tid der\b",
]


async def _ai_verdict(text: str) -> dict:
    """Ask Claude for a nuanced verdict on whether text feels AI-generated.
    Returns {label, confidence, reasoning} or {} if unavailable."""
    if not ANTHROPIC_API_KEY:
        return {}
    snippet = text[:6000]
    system = (
        "Du er en litterær kritiker med spesialkompetanse på å skille menneskelig prosa "
        "fra AI-generert tekst. Du kjenner igjen bevisste stilvalg — som Nick Hornbys "
        "samtale-aktige enkelthet eller Erlend Loes flate rytme — og forveksler dem "
        "IKKE med AI. Du er skeptisk til utjevnet, klisjéfull tekst med perfekt grammatikk "
        "og trippel-listestruktur. Svar KUN i gyldig JSON."
    )
    prompt = (
        "Vurder følgende norske tekst. Svar i JSON med:\n"
        "  - label: én av \"Menneskelig\", \"Sannsynligvis menneskelig\", \"Usikker\", \"Sannsynligvis AI\", \"AI-aktig\"\n"
        "  - confidence: heltall 0-100 (hvor sikker du er på vurderingen)\n"
        "  - reasoning: 1-2 setninger på norsk med konkret begrunnelse (nevn stiltrekk du observerer)\n"
        "  - notes: liste med 0-3 korte observasjoner om stemmen (fri form)\n\n"
        f"TEKST:\n{snippet}"
    )
    try:
        parts = []
        async for chunk in stream_llm_text(*BRAGR_MODEL_ANALYSIS, system, prompt, ANTHROPIC_API_KEY, temperature=0.3):
            parts.append(chunk)
        raw = "".join(parts).strip()
        m = re.search(r"\{.*\}", raw, re.DOTALL)
        if m:
            raw = m.group(0)
        return json.loads(raw)
    except Exception as e:
        logger.warning("AI verdict failed: %s", e)
        return {}


@api_router.post("/detect")
async def detect_ai(body: DetectBody, user: User = Depends(get_current_user)):
    text = body.text.strip()
    if not text:
        raise HTTPException(status_code=400, detail="Tom tekst")

    words = re.findall(r"[A-Za-zÆØÅæøå']+", text)
    total_words = max(len(words), 1)

    # Guard: heuristic + AI verdict need enough text to be meaningful
    MIN_WORDS_RELIABLE = 300
    too_short = total_words < MIN_WORDS_RELIABLE

    unique_ratio = len(set(w.lower() for w in words)) / total_words
    sentences = [s for s in re.split(r"(?<=[.!?])\s+", text) if s.strip()]
    sent_lens = [len(re.findall(r"\S+", s)) for s in sentences] or [0]
    if len(sent_lens) > 1:
        mean = sum(sent_lens) / len(sent_lens)
        variance = sum((x - mean) ** 2 for x in sent_lens) / len(sent_lens)
        std = variance ** 0.5
        burstiness = std / mean if mean else 0
    else:
        burstiness = 0

    marker_hits = 0
    hits_list = []
    for pat in AI_MARKERS:
        m = re.findall(pat, text, flags=re.IGNORECASE)
        if m:
            marker_hits += len(m)
            hits_list.append({"pattern": pat, "count": len(m)})

    # Heuristic score kept as SECONDARY signal — AI verdict is primary
    score = 0
    score += min(burstiness * 40, 40)
    score += min(unique_ratio * 60, 60)
    score -= marker_hits * 8
    score = max(0, min(100, score))
    heur_label = "Menneskelig" if score >= 65 else ("Blandet" if score >= 40 else "AI-aktig")

    # ---- AI verdict (only when text is long enough) ----
    verdict = {} if too_short else await _ai_verdict(text)

    # Final label: prefer AI verdict when available, otherwise heuristic
    label = verdict.get("label") or heur_label

    # ---- Personal style similarity ----
    personal = await compute_personal_style_score(user.user_id, text, words)

    highlights = []
    if personal.get("available"):
        profile = await db.voice_profiles.find_one({"user_id": user.user_id}, {"_id": 0})
        highlights = compute_sentence_highlights(text, profile)

    return {
        "score": round(score, 1),
        "label": label,
        "burstiness": round(burstiness, 3),
        "vocab_richness": round(unique_ratio, 3),
        "avg_sentence_length": round(sum(sent_lens) / len(sent_lens), 2) if sent_lens else 0,
        "ai_markers": hits_list,
        "personal_style": personal,
        "highlights": highlights,
        "too_short": too_short,
        "word_count": total_words,
        "min_words_reliable": MIN_WORDS_RELIABLE,
        "ai_verdict": verdict,   # {label, confidence, reasoning, notes} or {}
    }


def compute_sentence_highlights(text: str, profile: dict) -> List[dict]:
    """
    Score each sentence in the text against the user's voice profile.
    Returns list of {index, sentence, similarity (0..100), foreign (bool), foreign_words[]}.
    """
    if not profile:
        return []
    sentences = [s.strip() for s in re.split(r"(?<=[.!?])\s+", text.strip()) if s.strip()]
    if not sentences:
        return []

    # Precompute profile fingerprints
    profile_fw = {f["word"]: f["per_1000"] for f in (profile.get("function_word_frequencies") or [])}
    profile_top = set(w["word"] for w in (profile.get("top_words") or [])[:30])
    profile_avg_sl = profile.get("avg_sentence_length") or 0

    # Vocab from all profile samples: we approximate "known words" = signature + function words
    known_vocab = set(profile_fw.keys()) | profile_top

    ai_marker_res = [re.compile(pat, re.IGNORECASE) for pat in AI_MARKERS]

    results = []
    for idx, s in enumerate(sentences):
        s_words = re.findall(r"[A-Za-zÆØÅæøå']+", s)
        n = max(len(s_words), 1)
        s_lower = [w.lower() for w in s_words]

        # Function-word cosine (per-sentence, using absolute counts scaled)
        if profile_fw:
            counts = {}
            for w in s_lower:
                if w in profile_fw:
                    counts[w] = counts.get(w, 0) + 1
            per_1000 = 1000.0 / n
            s_fw = {w: c * per_1000 for w, c in counts.items()}
            keys = set(profile_fw.keys()) | set(s_fw.keys())
            dot = sum(profile_fw.get(k, 0) * s_fw.get(k, 0) for k in keys)
            na = sum(v * v for v in profile_fw.values()) ** 0.5
            nb = sum(v * v for v in s_fw.values()) ** 0.5
            fw_cos = (dot / (na * nb)) if (na and nb) else 0
        else:
            fw_cos = 0

        # Signature overlap
        overlap = sum(1 for w in s_lower if w in profile_top) / n

        # Sentence length delta
        if profile_avg_sl > 0:
            delta = abs(n - profile_avg_sl) / max(profile_avg_sl, 1)
            length_sim = max(0.0, 1.0 - min(delta, 1.0))
        else:
            length_sim = 0.5

        # AI marker in sentence => strong penalty
        ai_hit = any(rx.search(s) for rx in ai_marker_res)

        blended = (0.55 * fw_cos + 0.20 * overlap + 0.25 * length_sim) * 100
        if ai_hit:
            blended = max(0.0, blended - 25)
        blended = round(max(0.0, min(100.0, blended)), 1)

        # Foreign words in this sentence: content words not in known_vocab and appearing rarely
        foreign_words = []
        for w in s_words:
            wl = w.lower()
            if len(wl) < 5:
                continue
            if wl in NORWEGIAN_STOPWORDS or wl in known_vocab:
                continue
            if wl not in foreign_words:
                foreign_words.append(wl)
        foreign_words = foreign_words[:6]

        results.append({
            "index": idx,
            "sentence": s,
            "similarity": blended,
            "foreign": blended < 45,
            "foreign_words": foreign_words,
            "ai_marker_hit": ai_hit,
        })

    return results


async def compute_personal_style_score(user_id: str, text: str, words: List[str]) -> dict:
    """
    Similarity between the input text and the user's known voice.
    Uses:
      - Cosine similarity on function-word frequency vector
      - Overlap ratio of user's signature top-words appearing in the text
      - Sentence-length distribution correlation
    Returns a 0..100 personal_similarity score and diagnostic breakdown.
    """
    profile = await db.voice_profiles.find_one({"user_id": user_id}, {"_id": 0})
    if not profile:
        return {
            "available": False,
            "reason": "Ingen stemmeprofil ennå — kjør analyse under «Stemme» først.",
        }

    # Build vectors
    text_words = [w.lower() for w in words]
    text_len = max(len(text_words), 1)

    # Function-word cosine
    fw_profile = {f["word"]: f["per_1000"] for f in (profile.get("function_word_frequencies") or [])}
    if fw_profile:
        text_fw_counts = {}
        for w in text_words:
            if w in fw_profile:
                text_fw_counts[w] = text_fw_counts.get(w, 0) + 1
        per_1000 = 1000.0 / text_len
        text_fw = {w: c * per_1000 for w, c in text_fw_counts.items()}
        # Cosine over union of keys
        keys = set(fw_profile.keys()) | set(text_fw.keys())
        dot = sum(fw_profile.get(k, 0) * text_fw.get(k, 0) for k in keys)
        na = sum(v * v for v in fw_profile.values()) ** 0.5
        nb = sum(v * v for v in text_fw.values()) ** 0.5
        fw_cos = (dot / (na * nb)) if (na and nb) else 0
    else:
        fw_cos = 0

    # Signature-word overlap: how many of the user's top 15 words appear in this text
    top_words = [w["word"] for w in (profile.get("top_words") or [])[:15]]
    if top_words:
        text_word_set = set(text_words)
        overlap = sum(1 for w in top_words if w in text_word_set)
        overlap_ratio = overlap / len(top_words)
    else:
        overlap_ratio = 0

    # Sentence-length distribution correlation (pearson-ish, using shared bins)
    profile_dist = {d["range"]: d["count"] for d in (profile.get("sentence_length_distribution") or [])}
    if profile_dist:
        sentences = [s for s in re.split(r"(?<=[.!?])\s+", text) if s.strip()]
        sl = [len(re.findall(r"\S+", s)) for s in sentences]
        sl = [n for n in sl if n > 0]
        text_dist = {"1-5": 0, "6-10": 0, "11-15": 0, "16-20": 0, "21-30": 0, "31+": 0}
        for n in sl:
            if n <= 5: text_dist["1-5"] += 1
            elif n <= 10: text_dist["6-10"] += 1
            elif n <= 15: text_dist["11-15"] += 1
            elif n <= 20: text_dist["16-20"] += 1
            elif n <= 30: text_dist["21-30"] += 1
            else: text_dist["31+"] += 1
        # Normalize
        def norm(d):
            t = sum(d.values()) or 1
            return {k: v / t for k, v in d.items()}
        pn = norm(profile_dist)
        tn = norm(text_dist)
        # 1 - half L1 (0..1 similarity)
        dist_sim = 1 - 0.5 * sum(abs(pn[k] - tn[k]) for k in pn.keys())
    else:
        dist_sim = 0

    # Blend
    personal_similarity = (
        0.55 * fw_cos +
        0.25 * overlap_ratio +
        0.20 * dist_sim
    ) * 100
    personal_similarity = round(max(0, min(100, personal_similarity)), 1)

    if personal_similarity >= 70:
        label = "Din stemme"
    elif personal_similarity >= 45:
        label = "Delvis din stemme"
    else:
        label = "Fremmed stemme"

    return {
        "available": True,
        "personal_similarity": personal_similarity,
        "label": label,
        "function_word_cosine": round(fw_cos, 3),
        "signature_word_overlap": round(overlap_ratio, 3),
        "sentence_shape_similarity": round(dist_sim, 3),
    }


# ---------- Models list ----------
@api_router.get("/models")
async def list_models():
    """
    Public list of supported LLM models. Since Bragarmål picks the best model
    automatically for the user, this endpoint is mostly used internally / for
    admin/dev tooling. UI does not expose model selection anymore.
    """
    return [
        {"id": "claude-sonnet-4-5", "label": "Claude Sonnet 4.5", "provider": "Anthropic"},
        {"id": "claude-sonnet-4-6", "label": "Claude Sonnet 4.6 (Bragarmål-standard)", "provider": "Anthropic"},
        {"id": "claude-sonnet-5",   "label": "Claude Sonnet 5",   "provider": "Anthropic"},
        {"id": "claude-opus-4-7",   "label": "Claude Opus 4.7",   "provider": "Anthropic"},
        {"id": "claude-opus-5",     "label": "Claude Opus 5",     "provider": "Anthropic"},
        {"id": "gpt-5.2", "label": "GPT 5.2", "provider": "OpenAI"},
        {"id": "gpt-5.4", "label": "GPT 5.4", "provider": "OpenAI"},
        {"id": "gemini-3-pro", "label": "Gemini 3.1 Pro", "provider": "Google"},
        {"id": "gemini-3-flash", "label": "Gemini 3 Flash", "provider": "Google"},
    ]


@api_router.get("/files")
async def list_files(user: User = Depends(get_current_user)):
    docs = await db.files.find(
        {"user_id": user.user_id, "is_deleted": False}, {"_id": 0, "storage_path": 0}
    ).sort("created_at", -1).to_list(500)
    return docs


@api_router.get("/files/{file_id}/download")
async def download_file(
    file_id: str,
    authorization: Optional[str] = Header(None),
    auth: Optional[str] = Query(None),
    request: Request = None,
):
    # Auth: session_token cookie OR ?auth=token query param OR Authorization header
    token = None
    if request:
        token = request.cookies.get("session_token")
    if not token and auth:
        token = auth
    if not token and authorization and authorization.startswith("Bearer "):
        token = authorization.split(" ", 1)[1]
    if not token:
        raise HTTPException(status_code=401, detail="Ikke autentisert")

    session = await db.user_sessions.find_one({"session_token": token}, {"_id": 0})
    if not session:
        raise HTTPException(status_code=401, detail="Ugyldig økt")

    doc = await db.files.find_one(
        {"id": file_id, "user_id": session["user_id"], "is_deleted": False}, {"_id": 0}
    )
    if not doc:
        raise HTTPException(status_code=404, detail="Fil ikke funnet")

    data, content_type = storage_get(doc["storage_path"])
    return Response(
        content=data,
        media_type=doc.get("content_type") or content_type,
        headers={
            "Content-Disposition": f'inline; filename="{doc.get("original_filename", "fil")}"',
            "Cache-Control": "private, max-age=3600",
        },
    )


@api_router.delete("/files/{file_id}")
async def delete_file(file_id: str, user: User = Depends(get_current_user)):
    r = await db.files.update_one(
        {"id": file_id, "user_id": user.user_id, "is_deleted": False},
        {"$set": {"is_deleted": True, "deleted_at": datetime.now(timezone.utc).isoformat()}},
    )
    if r.matched_count == 0:
        raise HTTPException(status_code=404, detail="Ikke funnet")
    return {"ok": True}


async def get_user_rank(user_id: str) -> int:
    """1-based signup rank (1 = earliest)."""
    user = await db.users.find_one({"user_id": user_id}, {"_id": 0, "created_at": 1})
    if not user or not user.get("created_at"):
        return 999_999
    created = user["created_at"]
    cutoff = created if isinstance(created, str) else created.isoformat()
    earlier = await db.users.count_documents({"created_at": {"$lt": cutoff}})
    return earlier + 1


async def get_user_subscription_status(user_id: str) -> dict:
    """Return {active, plan, beta, current_period_end}"""
    now = datetime.now(timezone.utc)
    user = await db.users.find_one({"user_id": user_id}, {"_id": 0})

    # Lifetime free — configured emails never need to pay
    if user:
        lifetime_emails = {
            e.strip().lower()
            for e in os.environ.get("LIFETIME_FREE_EMAILS", "").split(",")
            if e.strip()
        }
        email = (user.get("email") or "").strip().lower()
        if email and email in lifetime_emails:
            return {"active": True, "plan": "lifetime", "beta": False, "current_period_end": None}

    if user and user.get("is_beta_member"):
        expires_at = user.get("beta_expires_at")
        if expires_at:
            try:
                exp = datetime.fromisoformat(expires_at) if isinstance(expires_at, str) else expires_at
                if exp.tzinfo is None:
                    exp = exp.replace(tzinfo=timezone.utc)
                if exp > now:
                    return {"active": True, "plan": "beta", "beta": True, "current_period_end": exp.isoformat()}
            except Exception:
                pass
        else:
            # Legacy beta without expiry — grant 3 months from creation
            return {"active": True, "plan": "beta", "beta": True, "current_period_end": None}

    # Any active subscription?
    sub = await db.subscriptions.find_one(
        {"user_id": user_id, "status": {"$in": ["active", "trialing"]}},
        {"_id": 0}, sort=[("current_period_end", -1)]
    )
    if not sub:
        return {"active": False, "plan": None, "beta": False, "current_period_end": None}
    return {
        "active": True,
        "plan": sub.get("lookup_key"),
        "beta": False,
        "current_period_end": sub.get("current_period_end"),
    }


async def ensure_beta_flag(user_id: str):
    """Grant beta ONLY to the first BETA_FREE_SLOTS users. Beta is free for 3 months."""
    user = await db.users.find_one({"user_id": user_id}, {"_id": 0})
    if not user or user.get("is_beta_member"):
        return
    rank = await get_user_rank(user_id)
    if rank <= BETA_FREE_SLOTS:
        expires_at = datetime.now(timezone.utc) + timedelta(days=90)
        await db.users.update_one(
            {"user_id": user_id},
            {"$set": {"is_beta_member": True, "beta_expires_at": expires_at.isoformat()}},
        )


def _ascii_url(url: str) -> str:
    """Stripe requires ASCII URLs — encode IDN hostnames (bragarmål.no → xn--bragarml-g0a.no)."""
    try:
        from urllib.parse import urlsplit, urlunsplit
        parts = urlsplit(url)
        host = parts.hostname or ""
        ascii_host = host.encode("idna").decode("ascii") if host else host
        if parts.port:
            ascii_host = f"{ascii_host}:{parts.port}"
        return urlunsplit((parts.scheme, ascii_host, parts.path, parts.query, parts.fragment))
    except Exception:
        return url


class CheckoutRequest(BaseModel):
    lookup_key: str
    origin_url: str
    trial_days: Optional[int] = None


@api_router.post("/billing/checkout")
async def billing_checkout(body: CheckoutRequest, user: User = Depends(get_current_user)):
    ALL_KEYS = {
        "bragr_monthly_nok", "bragr_3mo_nok", "bragr_6mo_nok", "bragr_yearly_nok",
        "bragr_monthly_founder", "bragr_yearly_founder",
    }
    if body.lookup_key not in ALL_KEYS:
        raise HTTPException(status_code=400, detail="Ukjent plan")

    # Enforce founder eligibility: only ranks 11-100 (or already founder subscribers) can buy founder price
    if "founder" in body.lookup_key:
        rank = await get_user_rank(user.user_id)
        existing_founder = await db.subscriptions.find_one({
            "user_id": user.user_id,
            "lookup_key": {"$in": [
                "bragr_monthly_founder", "bragr_yearly_founder",
            ]},
        }, {"_id": 0})
        if rank > FOUNDER_SLOTS and not existing_founder:
            raise HTTPException(
                status_code=403,
                detail="Grunnleggerprisen er kun for de 100 første brukerne. Velg vanlig pris.",
            )

    prices = stripe.Price.list(lookup_keys=[body.lookup_key], active=True, limit=1).data
    if not prices:
        raise HTTPException(status_code=500, detail="Pris ikke funnet")
    price = prices[0]

    # Find or create Stripe customer for this user
    user_doc = await db.users.find_one({"user_id": user.user_id}, {"_id": 0})
    customer_id = (user_doc or {}).get("stripe_customer_id")
    if not customer_id:
        cust = stripe.Customer.create(
            email=user.email,
            name=user.name,
            metadata={"user_id": user.user_id},
        )
        customer_id = cust.id
        await db.users.update_one({"user_id": user.user_id}, {"$set": {"stripe_customer_id": customer_id}})

    origin_ascii = _ascii_url(body.origin_url)

    # Build subscription_data with optional trial
    subscription_data = {"metadata": {"user_id": user.user_id, "lookup_key": body.lookup_key}}

    # Server-enforced 14-day trial for ordinary NOK plans, only for first-time subscribers
    NOK_PLANS = {"bragr_monthly_nok", "bragr_3mo_nok", "bragr_6mo_nok", "bragr_yearly_nok"}
    auto_trial_days = 0
    if body.lookup_key in NOK_PLANS:
        prior_sub = await db.subscriptions.find_one(
            {"user_id": user.user_id, "status": {"$in": ["active", "trialing", "past_due", "canceled"]}},
            {"_id": 0, "id": 1},
        )
        if not prior_sub:
            auto_trial_days = 14

    trial_days = body.trial_days if (body.trial_days and body.trial_days > 0) else auto_trial_days
    if trial_days and trial_days > 0:
        # Safety cap — trials over 60 days are almost certainly abuse
        subscription_data["trial_period_days"] = min(trial_days, 60)
        # Auto-charge with saved payment method after trial ends
        subscription_data["trial_settings"] = {
            "end_behavior": {"missing_payment_method": "cancel"}
        }

    session = stripe.checkout.Session.create(
        customer=customer_id,
        line_items=[{"price": price.id, "quantity": 1}],
        mode="subscription",
        success_url=f"{origin_ascii}/betaling/vellykket?session_id={{CHECKOUT_SESSION_ID}}",
        cancel_url=f"{origin_ascii}/betaling/avbrutt",
        metadata={"user_id": user.user_id, "lookup_key": body.lookup_key},
        managed_payments={"enabled": True},
        subscription_data=subscription_data,
        automatic_tax={"enabled": False},
        adaptive_pricing={"enabled": False},
    )

    await db.payment_transactions.insert_one({
        "session_id": session.id,
        "user_id": user.user_id,
        "lookup_key": body.lookup_key,
        "amount": price.unit_amount,
        "currency": price.currency,
        "status": "initiated",
        "payment_status": "pending",
        "created_at": datetime.now(timezone.utc).isoformat(),
        "updated_at": datetime.now(timezone.utc).isoformat(),
    })
    return {"checkout_url": session.url, "session_id": session.id}


@api_router.get("/billing/status")
async def billing_status(user: User = Depends(get_current_user)):
    await ensure_beta_flag(user.user_id)
    status = await get_user_subscription_status(user.user_id)
    rank = await get_user_rank(user.user_id)
    beta_users = await db.users.count_documents({"is_beta_member": True})
    beta_remaining = max(0, BETA_FREE_SLOTS - beta_users)
    founder_users = await db.users.count_documents({})
    founder_remaining = max(0, FOUNDER_SLOTS - founder_users)
    return {
        **status,
        "user_rank": rank,
        "beta_slots_remaining": beta_remaining,
        "beta_total": BETA_FREE_SLOTS,
        "founder_slots_remaining": founder_remaining,
        "founder_total": FOUNDER_SLOTS,
        "founder_eligible": rank <= FOUNDER_SLOTS,
    }


@api_router.get("/billing/session/{session_id}")
async def billing_session_status(session_id: str):
    record = await db.payment_transactions.find_one({"session_id": session_id}, {"_id": 0})
    if not record:
        raise HTTPException(status_code=404, detail="Betaling ikke funnet")

    if record.get("payment_status") != "paid":
        try:
            s = stripe.checkout.Session.retrieve(session_id)
            if s.payment_status == "paid" or s.status == "complete":
                await db.payment_transactions.update_one(
                    {"session_id": session_id, "payment_status": {"$ne": "paid"}},
                    {"$set": {
                        "status": "completed",
                        "payment_status": "paid",
                        "stripe_subscription_id": s.subscription,
                        "updated_at": datetime.now(timezone.utc).isoformat(),
                    }},
                )
                record = await db.payment_transactions.find_one({"session_id": session_id}, {"_id": 0})
        except Exception:
            pass

    return {
        "session_id": record["session_id"],
        "status": record["status"],
        "payment_status": record["payment_status"],
    }


class PortalRequest(BaseModel):
    return_url: str


@api_router.post("/billing/portal")
async def billing_portal(body: PortalRequest, user: User = Depends(get_current_user)):
    user_doc = await db.users.find_one({"user_id": user.user_id}, {"_id": 0})
    customer_id = (user_doc or {}).get("stripe_customer_id")
    if not customer_id:
        raise HTTPException(status_code=400, detail="Ingen kundeprofil ennå")
    portal = stripe.billing_portal.Session.create(
        customer=customer_id,
        return_url=body.return_url,
    )
    return {"portal_url": portal.url}


@api_router.post("/stripe/webhook")
async def stripe_webhook(request: Request):
    payload = await request.body()
    sig = request.headers.get("stripe-signature", "")
    try:
        event = stripe.Webhook.construct_event(payload, sig, STRIPE_WEBHOOK_SECRET)
    except Exception as e:
        logger.warning(f"Stripe webhook signature error: {e}")
        raise HTTPException(status_code=400, detail="Invalid signature")

    t = event["type"]
    obj = event["data"]["object"]
    now = datetime.now(timezone.utc).isoformat()

    if t == "checkout.session.completed":
        await db.payment_transactions.update_one(
            {"session_id": obj["id"], "payment_status": {"$ne": "paid"}},
            {"$set": {
                "status": "completed",
                "payment_status": obj.get("payment_status", "paid"),
                "stripe_subscription_id": obj.get("subscription"),
                "updated_at": now,
            }},
        )

    elif t in ("customer.subscription.created", "customer.subscription.updated"):
        illustrator_id = (obj.get("metadata") or {}).get("illustrator_id")
        if illustrator_id:
            await db.illustrators.update_one(
                {"id": illustrator_id},
                {"$set": {
                    "is_featured": obj.get("status") in ("active", "trialing"),
                    "stripe_subscription_id": obj["id"],
                    "updated_at": now,
                }},
            )
        else:
            user_id = (obj.get("metadata") or {}).get("user_id")
            lookup_key = (obj.get("metadata") or {}).get("lookup_key")
            if not user_id:
                # Look up by customer
                user_doc = await db.users.find_one({"stripe_customer_id": obj.get("customer")}, {"_id": 0})
                user_id = user_doc.get("user_id") if user_doc else None
            if user_id:
                await db.subscriptions.update_one(
                    {"stripe_subscription_id": obj["id"]},
                    {"$set": {
                        "user_id": user_id,
                        "stripe_subscription_id": obj["id"],
                        "stripe_customer_id": obj.get("customer"),
                        "status": obj.get("status"),
                        "lookup_key": lookup_key or "bragr_monthly_nok",
                        "current_period_end": datetime.fromtimestamp(obj["current_period_end"], tz=timezone.utc).isoformat() if obj.get("current_period_end") else None,
                        "cancel_at_period_end": obj.get("cancel_at_period_end", False),
                        "updated_at": now,
                    }},
                    upsert=True,
                )

    elif t == "customer.subscription.deleted":
        await db.subscriptions.update_one(
            {"stripe_subscription_id": obj["id"]},
            {"$set": {"status": "canceled", "updated_at": now}},
        )
        await db.illustrators.update_one(
            {"stripe_subscription_id": obj["id"]},
            {"$set": {"is_featured": False, "updated_at": now}},
        )

    return {"received": True}


@api_router.get("/")
async def root():
    return {"app": "BRAGARMÅL", "ok": True}


app.include_router(api_router)

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=os.environ.get('CORS_ORIGINS', '*').split(','),
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.on_event("shutdown")
async def shutdown_db_client():
    client.close()
